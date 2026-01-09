import streamlit as st
import google.generativeai as genai
import json
import re
import time
from PIL import Image
import random
import textwrap

# ==========================================
# 1. CẤU HÌNH API & PROMPT
# ==========================================

ALL_KEYS = st.secrets["GEMINI_API_KEYS"]

def generate_content_with_failover(prompt, image=None):
    """Hàm thông minh tự động dò tìm Model tốt nhất có sẵn lượt dùng"""
    keys_to_try = list(ALL_KEYS)
    random.shuffle(keys_to_try) 
    
    # DANH SÁCH ƯU TIÊN (Giữ nguyên theo yêu cầu của bạn)
    model_priority = [
        #"gemini-2.0-flash-thinking-preview-01-21",
        #"gemini-3-pro-preview", 
        #"gemini-2.5-pro",
        "gemini-3-flash-preview",        
        "gemini-2.5-flash",
        "gemini-2.5-flash-lite",
        "gemini-2.0-flash",
        "gemini-1.5-pro", 
        "gemini-1.5-flash"
    ]
    
    last_error = ""
    for index, current_key in enumerate(keys_to_try): # Thêm index để theo dõi số thứ tự Key
        try:
            genai.configure(api_key=current_key)
            
            # Lấy danh sách các model mà Key này THỰC SỰ có quyền truy cập
            available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            
            # Tìm model tốt nhất
            sel_model = None
            for target in model_priority:
                if any(target in m_name for m_name in available_models):
                    sel_model = target
                    break
            
            if not sel_model:
                sel_model = "gemini-1.5-flash" 

            # --- [PHẦN THÊM MỚI] HIỂN THỊ THÔNG TIN MODEL ---
            # 1. Tạo mặt nạ cho Key (Chỉ hiện 4 số cuối để bảo mật)
            masked_key = f"****{current_key[-4:]}"
            
            # 2. Hiển thị thông báo nhỏ góc màn hình (Toast)
            st.toast(f"⚡ Đã kết nối: {sel_model}", icon="🤖")
            
            # 3. Hiển thị chi tiết trong giao diện (Tùy chọn: Dùng Expander để không rối mắt)
            with st.expander("🔌 Technical Connection Details (Debug)", expanded=False):
                st.write(f"**Active Model:** `{sel_model}`")
                st.write(f"**Active API Key:** `{masked_key}` (Key #{index + 1})")
                if "thinking" in sel_model.lower():
                    st.caption("🧠 Thinking Mode: ON")
            # ------------------------------------------------
            
            # Khởi tạo model
            temp_model = genai.GenerativeModel(
                model_name=sel_model, 
                # system_instruction=GRADING_PROMPT_TEMPLATE # Bỏ comment dòng này nếu bạn đã define biến này ở ngoài
            )
            
            content_parts = [prompt]
            if image:
                content_parts.append(image)
                
             # Cấu hình Generation Config
            gen_config = {
                "temperature": 0.3,       # Tăng nhẹ từ 0.1 lên 0.3 để AI giải thích phong phú hơn
                "top_p": 0.95,            # Cho phép AI chọn từ vựng đa dạng hơn
                "top_k": 64,              # Giúp câu văn mượt mà và sâu sắc hơn
                "max_output_tokens": 32000, # Đảm bảo AI có đủ "đất" để viết bản phân tích dài
            }

            # Nếu là model Thinking thì thêm cấu hình suy luận
            if "thinking" in sel_model.lower():
                gen_config["thinking_config"] = {
                    "include_thoughts": True,
                    "thinking_budget": 32000 # Đẩy max ngân sách suy nghĩ
                }

            # Thực hiện gọi API
            response = temp_model.generate_content(
                content_parts,
                generation_config=gen_config
            )
            
            # Trả về kết quả và tên model để hiển thị thêm nếu cần
            return response, sel_model 
            
        except Exception as e:
            last_error = str(e)
            if "429" in last_error or "quota" in last_error.lower() or "limit" in last_error.lower():
                continue 
            else:
                break
                
    st.error(f"❌ Toàn bộ {len(keys_to_try)} Keys đã hết hạn mức. Lỗi cuối: {last_error}")
    return None, None 

# Prompt "khủng" từ file gemini.ts của bạn
GRADING_PROMPT_TEMPLATE = """
Bạn hãy đóng vai trò là "Examiner Tony". Nhiệm vụ của bạn là đọc bài viết dưới đây và tạo ra một báo cáo phân tích theo phong cách "Inline Annotation".

**YÊU CẦU ĐỊNH DẠNG OUTPUT (BẮT BUỘC):**
Kết quả trả về phải có 2 phần, ngăn cách bởi dấu "---":
1.  **Phần 1 (Bài viết đã highlight):** Copy lại toàn bộ bài viết của thí sinh. Với mỗi đoạn văn có lỗi hoặc điểm cần bình luận, hãy bọc đoạn đó trong thẻ `<span>` có class tương ứng (ví dụ: `<span class="comment-tg1">...</span>`).
2.  **Phần 2 (Danh sách bình luận):** Liệt kê các bình luận chi tiết, mỗi bình luận bắt đầu bằng `Commented [TG...]:`.

**VÍ DỤ OUTPUT MẪU:**
<span class="comment-tg1">The line graph gives information about...</span>
<span class="comment-tg2">Overall, it is clear that the amount of money...</span>
<span class="comment-tg3">Looking at the graph in more detail...</span>
---
**Commented [TG1]:** See LR.
**Commented [TG2]:** Clear overview. You could add that all asset classes except for company shares saw significant rises from around 2006.
**Commented [TG3]:** I'd group these two – both saw investment values fluctuate...

**QUY TẮC PHÂN TÍCH:**
*   **Tư duy "I'd do this":** Khi đưa ra gợi ý, hãy dùng văn phong cá nhân, ví dụ: *"I'd group these two..."* hoặc *"You could add that..."*.
*   **Tập trung vào Logic & Omission:** Ưu tiên bắt các lỗi về tư duy dữ liệu, cách nhóm thông tin, và sự thiếu sót chi tiết.
"""

# ==========================================
# 2. CẤU HÌNH GIAO DIỆN (Từ index.html & styles)
# ==========================================
st.set_page_config(page_title="IELTS Examiner Pro", page_icon="🛡️", layout="wide")

# CSS để tái tạo giao diện React (Màu sắc, Font Merriweather/Inter, Error Cards)
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Merriweather:wght@300;400;700&display=swap');
    
    /* Font và màu sắc Global */
    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
    }
    
    /* Style cho Header */
    h1 {
        font-family: 'Merriweather', serif !important;
        color: #0F172A !important;
        font-weight: 700 !important;
    }
    .pro-badge {
        color: #D40E14; 
        font-weight: bold;
    }
    .verified-badge {
        background-color: #F1F5F9;
        border: 1px solid #E2E8F0;
        padding: 4px 12px;
        border-radius: 99px;
        font-size: 14px;
        font-weight: bold;
        color: #475569;
        display: inline-flex;
        align-items: center;
        margin-left: 10px;
    }
    
    /* Style cho Error Cards (Giống MessageBubble.tsx) */
    .error-card {
        background-color: white;
        border: 1px solid #E5E7EB;
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 16px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
        transition: all 0.2s;
    }
    .error-card:hover {
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        border-color: #D1D5DB;
    }
    .error-header {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 12px;
        border-bottom: 1px solid #F3F4F6;
        padding-bottom: 8px;
    }
    .error-badge-grammar {
        background-color: #DCFCE7;
        border: 1px solid #22C55E;
        color: #022C22;
        padding: 2px 8px;
        border-radius: 6px;
        font-size: 14px;
        font-weight: 800;
        text-transform: uppercase;
    }
    .error-badge-vocab {
        background-color: #FEF9C3;
        border: 1px solid #FCD34D;
        color: #713F12;
        padding: 2px 8px;
        border-radius: 6px;
        font-size: 14px;
        font-weight: 800;
        text-transform: uppercase;
    }
    .impact-high { background-color: #FEE2E2; color: #991B1B; padding: 2px 8px; border-radius: 99px; font-size: 14px; font-weight: bold; border: 1px solid #FECACA;}
    .impact-medium { background-color: #FFEDD5; color: #9A3412; padding: 2px 8px; border-radius: 99px; font-size: 14px; font-weight: bold; border: 1px solid #FED7AA;}
    .impact-low { background-color: #DBEAFE; color: #1E40AF; padding: 2px 8px; border-radius: 99px; font-size: 14px; font-weight: bold; border: 1px solid #BFDBFE;}
    
    .correction-box {
        background-color: #F9FAFB;
        padding: 12px;
        border-radius: 8px;
        margin-bottom: 12px;
        font-size: 16px;
        border: 1px solid #F3F4F6;
    }
    
    /* Style cho Annotated Essay */
    .annotated-text {
        font-family: 'Merriweather', serif;
        line-height: 1.8;
        color: #374151;
        background-color: white;
        padding: 24px;
        border-radius: 12px;
        border: 1px solid #E5E7EB;
        box-shadow: 0 1px 2px rgba(0,0,0,0.05);
    }
    del {
        color: #9CA3AF;
        text-decoration: line-through;
        margin-right: 4px;
        text-decoration-thickness: 2px;
    }
    ins.grammar {
        background-color: #4ADE80;
        color: #022C22;
        text-decoration: none;
        padding: 2px 6px;
        border-radius: 4px;
        font-weight: 700;
        border: 1px solid #22C55E;
    }
    ins.vocab {
        background-color: #FDE047;
        color: #000;
        text-decoration: none;
        padding: 2px 6px;
        border-radius: 4px;
        font-weight: 700;
        border: 1px solid #FCD34D;
    }
    
    /* Button Style */
    div.stButton > button {
        background-color: #D40E14;
        color: white;
        font-weight: bold;
        border: none;
        padding: 10px 24px;
        border-radius: 8px;
        transition: all 0.3s;
    }
    div.stButton > button:hover {
        background-color: #B91C1C;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
    }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 3. LOGIC KẾT NỐI AI & XỬ LÝ DỮ LIỆU
# ==========================================

import html
import os
import requests
import re
import time
from io import BytesIO

# Thư viện Word
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Thư viện PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping

# --- HÀM TỰ ĐỘNG TẢI & ĐĂNG KÝ FONT ROBOTO (HỖ TRỢ TIẾNG VIỆT) ---
def clean_json(text):
    """Trích xuất JSON từ phản hồi của AI"""
    match = re.search(r"```json\s*([\s\S]*?)\s*```", text)
    if match:
        content = match.group(1)
        content = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', content)
        return content.strip()
    return None

def calculate_overall(scores):
    """Hàm phụ: Tự tính điểm Overall chuẩn IELTS từ 4 điểm thành phần"""
    try:
        # Lọc bỏ các giá trị không phải số
        valid_scores = []
        for s in scores:
            try:
                valid_scores.append(float(s))
            except:
                continue
                
        if not valid_scores or len(valid_scores) < 4: return '-'
        
        avg = sum(valid_scores) / len(valid_scores)
        decimal = avg - int(avg)
        
        # Quy tắc làm tròn IELTS (.25 lên .5, .75 lên 1.0)
        if decimal < 0.25: final = int(avg)
        elif decimal < 0.75: final = int(avg) + 0.5
        else: final = int(avg) + 1.0
        
        return str(final)
    except:
        return '-'

def process_response(full_text):
    """
    1. Tách JSON để lấy danh sách lỗi và bài sửa.
    2. Dùng Regex quét văn bản Markdown để lấy điểm Gốc (Original Score).
    """
    json_str = clean_json(full_text)
    markdown_part = full_text
    
    # Khởi tạo cấu trúc dữ liệu mặc định
    data = {
        "errors": [], 
        "annotatedEssay": None, 
        "revisedScore": None, 
        "originalScore": {
            "task_achievement": "-",
            "cohesion_coherence": "-",
            "lexical_resource": "-",
            "grammatical_range": "-",
            "overall": "-"
        }
    }
    
    # A. Xử lý JSON (Chủ yếu để lấy Lỗi và Bài sửa)
    if json_str:
        markdown_part = full_text.split("```json")[0].strip()
        try:
            parsed = json.loads(json_str)
            data["errors"] = parsed.get("errors", [])
            data["annotatedEssay"] = parsed.get("annotated_essay")
            data["revisedScore"] = parsed.get("revised_score")
        except json.JSONDecodeError:
            pass

    # B. QUAN TRỌNG: COPY ĐIỂM TỪ VĂN BẢN (Regex Scanning)
    patterns = {
        "task_achievement": r"Điểm\s+Task\s+Achievement.*?(\d+\.?\d*)",
        "cohesion_coherence": r"Điểm\s+Coherence.*?(\d+\.?\d*)",
        "lexical_resource": r"Điểm\s+Lexical.*?(\d+\.?\d*)",
        "grammatical_range": r"Điểm\s+Grammatical.*?(\d+\.?\d*)",
    }
    
    found_scores = []
    
    # Quét từng tiêu chí trong văn bản markdown
    for key, regex in patterns.items():
        match = re.search(regex, markdown_part, re.IGNORECASE | re.DOTALL)
        if match:
            score = match.group(1) # Lấy con số tìm được
            data["originalScore"][key] = score
            found_scores.append(score)
        else:
            # Nếu không tìm thấy, thử tìm trong JSON dự phòng
            try:
                if json_str:
                    parsed = json.loads(json_str)
                    val = parsed.get("original_score", {}).get(key, "-")
                    data["originalScore"][key] = str(val)
                    if str(val) != "-": found_scores.append(val)
            except:
                pass

    # C. Tự tính Overall từ các điểm vừa tìm được
    if found_scores:
        data["originalScore"]["overall"] = calculate_overall(found_scores)

    return markdown_part, data

# --- 2. CÁC HÀM XUẤT FILE (FILE EXPORT) ---

def register_vietnamese_font():
    """Tải và đăng ký font Roboto từ Google Fonts (Github Source)"""
    font_reg = "Roboto-Regular.ttf"
    font_bold = "Roboto-Bold.ttf"
    
    # URL mới chính xác (Dùng kho 'src/hinted' của googlefonts)
    urls = {
        font_reg: "https://github.com/googlefonts/roboto/raw/main/src/hinted/Roboto-Regular.ttf",
        font_bold: "https://github.com/googlefonts/roboto/raw/main/src/hinted/Roboto-Bold.ttf"
    }
    
    # Headers giả lập trình duyệt để tránh lỗi 403/404
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }

    try:
        # 1. Tải font về nếu chưa có hoặc file bị lỗi (0kb)
        for filename, url in urls.items():
            if os.path.exists(filename) and os.path.getsize(filename) < 1000:
                os.remove(filename) # Xóa file lỗi
                
            if not os.path.exists(filename):
                response = requests.get(url, headers=headers, timeout=20)
                if response.status_code == 200:
                    with open(filename, "wb") as f:
                        f.write(response.content)
                else:
                    print(f"❌ Không tải được {filename}. Mã lỗi: {response.status_code}")

        # 2. Đăng ký với ReportLab
        if os.path.exists(font_reg) and os.path.exists(font_bold):
            pdfmetrics.registerFont(TTFont('Roboto', font_reg))
            pdfmetrics.registerFont(TTFont('Roboto-Bold', font_bold))
            addMapping('Roboto', 0, 0, 'Roboto') # Normal
            addMapping('Roboto', 1, 0, 'Roboto-Bold') # Bold
            return True
        else:
            print("⚠️ Không tìm thấy file font sau khi tải.")
            return False
            
    except Exception as e:
        print(f"❌ Lỗi xử lý font: {e}")
        return False

# HÀM TẠO FILE WORD
def create_docx(data, topic, original_essay, analysis_text):
    doc = Document()
    
    # Header
    heading = doc.add_heading('IELTS WRITING TASK 1 - ASSESSMENT REPORT', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date: {time.strftime('%d/%m/%Y')}")
    
    # 1. BAND SCORE
    doc.add_heading('1. BAND SCORE', level=1)
    scores = data.get("originalScore")
    
    if scores and isinstance(scores, dict) and scores.get('overall', '-') != '-':
        table = doc.add_table(rows=2, cols=5)
        table.style = 'Table Grid'
        
        headers = ['Task Achievement', 'Coherence', 'Lexical Resource', 'Grammar', 'OVERALL']
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
        
        vals = [
            str(scores.get('task_achievement', '-')),
            str(scores.get('cohesion_coherence', '-')),
            str(scores.get('lexical_resource', '-')),
            str(scores.get('grammatical_range', '-')),
            str(scores.get('overall', '-'))
        ]
        for i, v in enumerate(vals):
            table.cell(1, i).text = v
    else:
        doc.add_paragraph("Score details could not be extracted automatically.")

    # 2. ANALYSIS
    doc.add_heading('2. EXAMINER\'S DETAILED ANALYSIS', level=1)
    if analysis_text:
        clean_analysis = analysis_text.replace('**', '').replace('### ', '').replace('#### ', '')
        doc.add_paragraph(clean_analysis)

    # 3. ERRORS
    doc.add_heading('3. DETAILED ERROR LOG', level=1)
    if data.get("errors"):
        for err in data['errors']:
            p = doc.add_paragraph(style='List Bullet')
            runner = p.add_run(f"[{err['category']} - {err['type']}]: ")
            runner.bold = True
            runner.font.color.rgb = RGBColor(200, 0, 0)
            p.add_run(f" '{err['original']}' → '{err['correction']}'")
            p.add_run(f"\n   Reason: {err['explanation']}")
    else:
        doc.add_paragraph("No specific errors detected.")

    # APPENDIX
    doc.add_page_break()
    doc.add_heading('APPENDIX', level=1)
    doc.add_heading('A. Task Prompt:', level=2)
    doc.add_paragraph(topic)
    doc.add_heading('B. Original Essay:', level=2)
    doc.add_paragraph(original_essay)
    doc.add_heading('C. Annotated Version:', level=2)
    clean_annotated = re.sub(r'<[^>]+>', '', data.get("annotatedEssay", "") or "")
    doc.add_paragraph(clean_annotated)

    # D. PROJECTED SCORE
    doc.add_heading('D. PROJECTED BAND SCORE (AFTER REVISION)', level=2)
    rev_scores = data.get("revisedScore")
    if rev_scores:
        table = doc.add_table(rows=2, cols=5)
        table.style = 'Table Grid'
        vals = [
            str(rev_scores.get('task_achievement', '-')),
            str(rev_scores.get('cohesion_coherence', '-')),
            str(rev_scores.get('lexical_resource', '-')),
            str(rev_scores.get('grammatical_range', '-')),
            str(rev_scores.get('overall', '-'))
        ]
        # Header
        for i, h in enumerate(['Task Achievement', 'Coherence', 'Lexical Resource', 'Grammar', 'OVERALL']):
            cell = table.cell(0, i)
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 100, 0)
        # Value
        for i, v in enumerate(vals):
            table.cell(1, i).text = v
            
        if rev_scores.get('logic_re_evaluation'):
            p = doc.add_paragraph()
            run = p.add_run(f"\nExaminer's Note: {rev_scores['logic_re_evaluation']}")
            run.font.italic = True
            run.font.color.rgb = RGBColor(0, 128, 0)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# HÀM TẠO FILE PDF
def create_pdf(data, topic, original_essay, analysis_text):
    # 1. Đăng ký Font
    has_font = register_vietnamese_font()
    font_name = 'Roboto' if has_font else 'Helvetica'
    font_bold = 'Roboto-Bold' if has_font else 'Helvetica-Bold'

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    
    # 2. Lấy Styles
    styles = getSampleStyleSheet()
    
    # 3. Cập nhật Font
    styles['Title'].fontName = font_name
    styles['Title'].fontSize = 18
    styles['Heading1'].fontName = font_bold
    styles['Heading2'].fontName = font_bold
    styles['Normal'].fontName = font_name
    styles['Normal'].fontSize = 13
    
    h1_style = styles['Heading1']
    h2_style = styles['Heading2']
    normal_style = styles['Normal']
    
    elements = []

    # Title
    elements.append(Paragraph("IELTS WRITING ASSESSMENT REPORT", styles['Title']))
    elements.append(Spacer(1, 12))

    # 1. BAND SCORE
    elements.append(Paragraph("1. BAND SCORE", h1_style))
    scores = data.get("originalScore")
    
    if scores and isinstance(scores, dict) and scores.get('overall', '-') != '-':
        data_table = [
            ['TA', 'CC', 'LR', 'GRA', 'OVERALL'],
            [
                str(scores.get('task_achievement', '-')),
                str(scores.get('cohesion_coherence', '-')),
                str(scores.get('lexical_resource', '-')),
                str(scores.get('grammatical_range', '-')),
                str(scores.get('overall', '-'))
            ]
        ]
        t = Table(data_table, colWidths=[60, 60, 60, 60, 80])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkred),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 0), (-1, -1), font_name) 
        ]))
        elements.append(t)
    else:
        elements.append(Paragraph("Original score data not found.", normal_style))
    
    elements.append(Spacer(1, 12))

    # 2. ANALYSIS
    elements.append(Paragraph("2. DETAILED ANALYSIS", h1_style))
    if analysis_text:
        safe_text = html.escape(analysis_text).replace('\n', '<br/>').replace('**', '').replace('###', '')
        elements.append(Paragraph(safe_text, normal_style))
    else:
        elements.append(Paragraph("No detailed analysis available.", normal_style))
    elements.append(Spacer(1, 12))

    # 3. ERRORS
    elements.append(Paragraph("3. ERROR LOG", h1_style))
    if data.get("errors"):
        for err in data['errors']:
            cat = html.escape(str(err.get('category', '')))
            typ = html.escape(str(err.get('type', '')))
            orig = html.escape(str(err.get('original', '')))
            fix = html.escape(str(err.get('correction', '')))
            text = f"<b>[{cat}] {typ}</b><br/>Original: <strike>{orig}</strike> -> Fix: <b>{fix}</b>"
            elements.append(Paragraph(text, normal_style))
            elements.append(Spacer(1, 6))

    # APPENDIX
    elements.append(PageBreak())
    elements.append(Paragraph("APPENDIX", h1_style))
    
    elements.append(Paragraph("<b>A. Task Prompt:</b>", h2_style))
    elements.append(Paragraph(html.escape(topic).replace('\n', '<br/>'), normal_style))
    elements.append(Spacer(1, 10))
    
    elements.append(Paragraph("<b>B. Original Essay:</b>", h2_style))
    elements.append(Paragraph(html.escape(original_essay).replace('\n', '<br/>'), normal_style))
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("<b>C. Annotated Version:</b>", h2_style))
    clean_annotated = re.sub(r'<[^>]+>', '', data.get("annotatedEssay", "") or "")
    elements.append(Paragraph(html.escape(clean_annotated).replace('\n', '<br/>'), normal_style))
    elements.append(Spacer(1, 10))

    # D. PROJECTED
    elements.append(Paragraph("<b>D. PROJECTED BAND SCORE (AFTER REVISION):</b>", h2_style))
    rev_scores = data.get("revisedScore")
    if rev_scores:
        rev_table_data = [
            ['TA', 'CC', 'LR', 'GRA', 'OVERALL'],
            [
                str(rev_scores.get('task_achievement', '-')),
                str(rev_scores.get('cohesion_coherence', '-')),
                str(rev_scores.get('lexical_resource', '-')),
                str(rev_scores.get('grammatical_range', '-')),
                str(rev_scores.get('overall', '-'))
            ]
        ]
        t2 = Table(rev_table_data, colWidths=[60, 60, 60, 60, 80])
        t2.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkgreen),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 0), (-1, -1), font_name)
        ]))
        elements.append(t2)
        
        if rev_scores.get('logic_re_evaluation'):
            safe_note = html.escape(rev_scores['logic_re_evaluation'])
            elements.append(Spacer(1, 6))
            elements.append(Paragraph(f"<i>Examiner's Note: {safe_note}</i>", normal_style))

    doc.build(elements)
    buffer.seek(0)
    return buffer
    
# ==========================================
# 4. GIAO DIỆN CHÍNH (UI)
# ==========================================
import textwrap

# HEADER (Thanh tiêu đề và nút xóa)
c1, c2 = st.columns([3, 1])
with c1:
    st.markdown("""
        <div style="display: flex; flex-direction: column; justify-content: center;">
            <h1 style="margin-bottom: 5px; line-height: 0.2rem;">
                IELTS Examiner <span class='pro-badge'>Pro</span>
            </h1>
            <div>
                <span class='verified-badge' style="margin-left: 2px;">
                    🛡️ BC CERTIFIED EXPERT
                </span>
            </div>
        </div>
    """, unsafe_allow_html=True)
with c2:
    if st.button("🗑️ Clear Session", use_container_width=True):
        st.session_state.messages = []
        st.session_state.submitted = False 
        st.rerun()

if "submitted" not in st.session_state:
    st.session_state.submitted = False

if "messages" not in st.session_state:
    st.session_state.messages = [
        {
            "role": "ai",
            "content": """
<div style="font-family: 'Inter', sans-serif; color: #1F2937; line-height: 1.4; font-size: 16px; max-width: 850px;">
    <h3 style="color: #D40E14; font-family: 'Merriweather', serif; margin-top: 0; margin-bottom: 15px; font-size: 22px; border-bottom: 3px solid #D40E14; display: inline-block; padding-bottom: 5px;">
        Welcome to the Official Task 1 Assessment Room.
    </h3>
    <p style="margin-bottom: 13px;">
        This system provides expert-level evaluation of <b>IELTS Academic Task 1 reports</b>, based on the official IELTS band descriptors.
    </p>
    <p style="margin-bottom: 13px;">
        The assessment focuses on objective, criteria-based feedback to help you understand your current writing level and areas for improvement.
    </p>
    <div style="background-color: #F8FAFC; border-radius: 8px; padding: 15px 20px; border-left: 5px solid #D40E14; box-shadow: 0 2px 4px rgba(0,0,0,0.05);">
        <p style="margin-bottom: 10px; font-weight: 800; font-size: 13px; color: #111827; text-transform: uppercase; letter-spacing: 1px;">
            Guidelines for a valid submission:
        </p>
        <div style="color: #374151;">
            <div style="margin-bottom: 6px;">• <b>Task Prompt:</b> Provide the original question or instruction.</div>
            <div style="margin-bottom: 6px;">• <b>Visual Data:</b> Upload a clear image of the chart, graph, table, or diagram.</div>
            <div>• <b>Written Report:</b> Paste your complete response (at least <b>150 words</b> to avoid penalties).</div>
        </div>
    </div>
</div>
""",
            "data": None
        }
    ]

# HIỂN THỊ LỊCH SỬ CHAT VÀ KẾT QUẢ CHẤM ĐIỂM
for msg in st.session_state.messages:
    with st.chat_message(msg["role"], avatar="👨‍🏫" if msg["role"] == "ai" else "👤"):
        if msg["role"] == "user":
            if msg.get("topic"):
                st.markdown(f"**📝 Task Prompt:**\n> {msg['topic']}")
            if msg.get("image"):
                st.image(msg["image"], caption="Visual Resource Attached", width=400)
            st.write(msg["content"])
        else:
            # --- PHẦN HIỂN THỊ MỚI THEO PHONG CÁCH "COMMENT" ---
            
            # 1. TÁCH BÀI VIẾT VÀ BÌNH LUẬN TỪ KẾT QUẢ CỦA AI
            content_parts = msg["content"].split("\n---\n", 1)
            essay_html = content_parts[0]
            comments_text = content_parts[1] if len(content_parts) > 1 else ""

            # 2. PARSE CÁC BÌNH LUẬN
            comments_dict = {}
            # Dùng Regex để tách từng comment ra (TG1, TG2, ...)
            matches = re.findall(r"Commented\s*\[(TG\d+)\]:\s*([\s\S]*?)(?=\nCommented|\Z)", comments_text)
            for match in matches:
                tag_id = match[0].lower() # -> tg1, tg2
                comment_content = match[1].strip()
                comments_dict[tag_id] = comment_content

            # 3. TẠO HTML HOÀN CHỈNH VỚI TOOLTIP
            # Thêm CSS cho tooltip
            st.markdown("""
            <style>
                .comment-highlight {
                    background-color: #fce7f3; /* Màu hồng nhạt */
                    border-bottom: 2px dashed #db2777;
                    cursor: pointer;
                    position: relative;
                }
                .tooltip {
                    visibility: hidden;
                    width: 250px;
                    background-color: #3B82F6;
                    color: #fff;
                    text-align: left;
                    border-radius: 6px;
                    padding: 10px;
                    position: absolute;
                    z-index: 1;
                    bottom: 125%;
                    left: 50%;
                    margin-left: -125px;
                    opacity: 0;
                    transition: opacity 0.3s;
                }
                .comment-highlight:hover .tooltip {
                    visibility: visible;
                    opacity: 1;
                }
            </style>
            """, unsafe_allow_html=True)
            
            # Thay thế các class trong HTML bằng thẻ có tooltip
            final_html = essay_html
            for tag_id, comment in comments_dict.items():
                replacement_html = f'''
                <span class="comment-highlight">
                    {tag_id.upper()}
                    <span class="tooltip">{comment}</span>
                </span>
                '''
                # Tìm và thay thế thẻ span tương ứng
                final_html = re.sub(f'<span class="comment-{tag_id}">', f'<span class="comment-highlight">', final_html, flags=re.IGNORECASE)
                # Thay thế thẻ đóng, đồng thời chèn tooltip vào
                final_html = re.sub(r'</span>', f'{replacement_html}</span>', final_html, count=1)


            # 4. HIỂN THỊ GIAO DIỆN
            st.subheader("📝 Examiner's Review")
            st.markdown(final_html, unsafe_allow_html=True)
            
            # --- CÁC PHẦN CÒN LẠI (BẢNG ĐIỂM, DOWNLOAD...) ---
            # Bạn có thể giữ lại phần hiển thị Bảng điểm, nút Download... ở dưới đây
            # Ví dụ:
            if msg.get("data") and msg["data"].get("revisedScore"):
                 scores = msg["data"]["revisedScore"]
                 st.markdown("### 📊 Band Scores")
                 cols = st.columns(5)
                 # ... (code hiển thị điểm)
# ==========================================
# 5. KHU VỰC NHẬP LIỆU & XỬ LÝ AI (ẨN KHI XONG)
# ==========================================

# Chỉ hiển thị nếu chưa nộp bài
if not st.session_state.submitted:
    st.markdown("---")
    with st.container():
        # Bố cục bất đối xứng: Cột trái 1.3 - Cột phải 2.7
        col_left, col_right = st.columns([1.3, 2.7], gap="large")
        
        with col_left:
            st.markdown("<p style='font-weight: 700; font-size: 15px; color: #1F2937;'>📝 TASK 1 QUESTION / PROMPT</p>", unsafe_allow_html=True)
            topic_input = st.text_area("topic_label", label_visibility="collapsed", height=280, placeholder="Paste the official question text here...")
            
            st.markdown("<div style='margin-top: 25px;'></div>", unsafe_allow_html=True)
            st.markdown("<p style='font-weight: 700; font-size: 15px; color: #1F2937;'>📊 VISUAL DATA</p>", unsafe_allow_html=True)
            uploaded_file = st.file_uploader("file_label", label_visibility="collapsed", type=['png', 'jpg', 'jpeg'])
            
        with col_right:
            st.markdown("<p style='font-weight: 700; font-size: 15px; color: #1F2937;'>✍️ YOUR WRITTEN REPORT</p>", unsafe_allow_html=True)
            essay_input = st.text_area("essay_label", label_visibility="collapsed", height=515, placeholder="Type or paste your response here (aim for 150+ words)...")

        st.markdown("<div style='margin-top: 20px;'></div>", unsafe_allow_html=True)
        submit_btn = st.button("🚀 SUBMIT FOR ASSESSMENT", type="primary", use_container_width=True)

        if submit_btn:
            # KIỂM TRA BẮT BUỘC NHẬP ĐỦ 3 THÔNG TIN
            if not topic_input.strip():
                st.warning("⚠️ Bắt buộc: Vui lòng nhập đề bài (Task Prompt) trước khi chấm điểm!")
            elif uploaded_file is None:
                st.warning("⚠️ Bắt buộc: Vui lòng tải lên hình ảnh biểu đồ (Visual Data) để Giám khảo đối chiếu số liệu!")
            elif not essay_input.strip() or len(essay_input.strip()) < 10:
                st.warning("⚠️ Bắt buộc: Vui lòng nhập nội dung bài làm (ít nhất 10 ký tự)!")
            else:
                # NẾU ĐÃ NHẬP ĐỦ, BẮT ĐẦU QUY TRÌNH CHẤM ĐIỂM
                loading_steps = [
                    "🕵️ INITIAL VALIDATION: IDENTIFYING EXAM CONTEXT AND ENFORCING WORD COUNT CONSTRAINTS...",
                    "🔍 EXHAUSTIVE ERROR SCANNING: CONDUCTING SENTENCE-BY-SENTENCE REVIEW FOR ALL ERRORS...",
                    "📊 DEEP CRITERIA ANALYSIS: EVALUATING TA, CC, LR, AND GRA STANDARDS WITH CEILING SCORES...",
                    "🧮 SCORE CALCULATION: DETERMINING COMPONENT BANDS AND APPLYING IELTS ROUNDING RULES...",
                    "⚖️ CONSISTENCY CHECK: CROSS-REFERENCING ASSIGNED SCORES WITH ERROR LOG FOR LOGICAL ACCURACY...",
                    "📝 OUTPUT GENERATION: COMPILING DETAILED ANALYSIS AND PRODUCING ANNOTATED ESSAY DATA..."
                ]
                
                status_container = st.status("👨‍🏫 Senior Examiner is starting assessment...", expanded=True)
                progress_bar = status_container.progress(0)
                
                try:
                    # 1. Xử lý hình ảnh
                    image_part = Image.open(uploaded_file)
                    
                    # 2. Ghép Prompt
                    full_prompt = GRADING_PROMPT_TEMPLATE.replace('{{TOPIC}}', topic_input).replace('{{ESSAY}}', essay_input)
                    
                    # 3. Gọi AI
                    response, used_model = generate_content_with_failover(full_prompt, image_part)
                    
                    # 4. Chạy hiệu ứng Loading
                    for i, text in enumerate(loading_steps):
                        status_container.write(text)
                        progress_bar.progress(int((i + 1) * (100 / len(loading_steps))))
                        time.sleep(2.8) 
                    
                    if response:
                        markdown_text, parsed_data = process_response(response.text)
                        st.session_state.messages.append({"role": "user", "content": essay_input, "topic": topic_input, "image": uploaded_file})
                        st.session_state.messages.append({"role": "ai", "content": markdown_text, "data": parsed_data, "model_version": used_model})
                        st.session_state.submitted = True
                        status_container.update(label="✅ ASSESSMENT COMPLETE!", state="complete", expanded=False)
                        st.rerun()
                        
                except Exception as e:
                    status_container.update(label="❌ Error occurred!", state="error")
                    st.error(f"Lỗi hệ thống: {e}")
# Footer
st.markdown("---")
st.caption("Developed by Albert Nguyen - v20251228.")




















