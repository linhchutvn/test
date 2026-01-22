import streamlit as st
from google import genai
from google.genai import types
import json
import re
import time
import random
import textwrap
import html
import os
import requests
from PIL import Image
from io import BytesIO

# Thư viện Word
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Thư viện PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping

# ==========================================
# 1. CẤU HÌNH TRANG (PHẢI ĐẶT ĐẦU TIÊN)
# ==========================================
st.set_page_config(page_title="IELTS Writing Master", page_icon="🎓", layout="wide")

# ==========================================
# 2. CSS TỔNG HỢP (ẨN HEADER/FOOTER + STYLE APP)
# ==========================================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Merriweather:wght@300;400;700&display=swap');
    
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

    /* --- PHẦN ẨN GIAO DIỆN MẶC ĐỊNH --- */
    
    /* 1. Ẩn thanh Header trên cùng (Chứa nút 3 chấm và Running man) */
    .stAppHeader {
        display: none;
    }
    
    /* 2. Ẩn Footer 'Made with Streamlit' */
    footer {
        visibility: hidden;
    }
    
    /* 3. Ẩn nút Deploy (Con thuyền màu đỏ) */
    .stDeployButton {
        display: none;
    }
    
    /* 4. Ẩn Menu Hamburger (nếu CSS trên chưa ẩn hết) */
    #MainMenu {
        visibility: hidden;
    }

    /* --- PHẦN STYLE GIAO DIỆN APP --- */
    
    /* Header Style */
    .main-header {
        font-family: 'Merriweather', serif;
        color: #0F172A;
        font-weight: 700;
        font-size: 2.2rem;
        margin-bottom: 0rem;
        margin-top: -2rem; /* Đẩy tiêu đề lên cao hơn vì đã ẩn Header */
    }
    .sub-header {
        font-family: 'Inter', sans-serif;
        color: #64748B;
        font-size: 1.1rem;
        margin-bottom: 1rem;
        border-bottom: 1px solid #E2E8F0;
        padding-bottom: 0.5rem;
    }

    /* Step Headers */
    .step-header {
        font-family: 'Inter', sans-serif;
        font-weight: 700;
        font-size: 1.2rem;
        color: #1E293B;
        margin-top: 1.5rem;
        margin-bottom: 0.5rem;
    }
    .step-desc {
        font-size: 0.9rem;
        color: #64748B;
        margin-bottom: 0.8rem;
    }
    /* --- ẨN CÁC ICON GHIM (LINK CHAIN) BÊN CẠNH TIÊU ĐỀ --- */
    [data-testid="stMarkdownContainer"] h1 a,
    [data-testid="stMarkdownContainer"] h2 a,
    [data-testid="stMarkdownContainer"] h3 a,
    [data-testid="stMarkdownContainer"] h4 a,
    [data-testid="stMarkdownContainer"] h5 a,
    [data-testid="stMarkdownContainer"] h6 a {
        display: none !important;
        pointer-events: none;
    }

    /* Guide Box */
    .guide-box {
        background-color: #f8f9fa;
        border-left: 5px solid #ff4b4b;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 10px;
        color: #31333F;
    }

    /* Error Cards */
    .error-card {
        background-color: white;
        border: 1px solid #E5E7EB;
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 16px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
        transition: all 0.2s;
    }
    .error-card:hover {
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        border-color: #D1D5DB;
    }
    
    .annotated-text {
        font-family: 'Merriweather', serif;
        line-height: 1.8;
        color: #374151;
        background-color: white;
        padding: 24px;
        border-radius: 12px;
        border: 1px solid #E5E7EB;
        box-shadow: 0 1px 2px rgba(0,0,0,0.05);
    }
    
    del { color: #9CA3AF; text-decoration: line-through; margin-right: 4px; text-decoration-thickness: 2px; }
    ins.grammar { background-color: #4ADE80; color: #022C22; text-decoration: none; padding: 2px 6px; border-radius: 4px; font-weight: 700; border: 1px solid #22C55E; }
    ins.vocab { background-color: #FDE047; color: #000; text-decoration: none; padding: 2px 6px; border-radius: 4px; font-weight: 700; border: 1px solid #FCD34D; }
    
    /* Button Customization */
    div.stButton > button {
        background-color: #FF4B4B;
        color: white;
        font-weight: bold;
        border-radius: 8px;
        padding: 0.5rem 1.5rem;
        border: none;
    }
    div.stButton > button:hover {
        background-color: #D93434;
    }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 2. LOGIC AI (FAILOVER)
# ==========================================
try:
    ALL_KEYS = st.secrets["GEMINI_API_KEYS"]
except Exception:
    st.error("⚠️ Chưa cấu hình secrets.toml chứa GEMINI_API_KEYS!")
    st.stop()

import streamlit as st
from google import genai
from google.genai import types
import json
import re
import time
import random
import textwrap
import html
import os
import requests
from PIL import Image
from io import BytesIO

# Thư viện Word
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Thư viện PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping

# ==========================================
# 1. CẤU HÌNH TRANG (PHẢI ĐẶT ĐẦU TIÊN)
# ==========================================
st.set_page_config(page_title="IELTS Writing Master", page_icon="🎓", layout="wide")

# ==========================================
# 2. CSS TỔNG HỢP (ẨN HEADER/FOOTER + STYLE APP)
# ==========================================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Merriweather:wght@300;400;700&display=swap');
    
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

    /* --- PHẦN ẨN GIAO DIỆN MẶC ĐỊNH --- */
    
    /* 1. Ẩn thanh Header trên cùng (Chứa nút 3 chấm và Running man) */
    .stAppHeader {
        display: none;
    }
    
    /* 2. Ẩn Footer 'Made with Streamlit' */
    footer {
        visibility: hidden;
    }
    
    /* 3. Ẩn nút Deploy (Con thuyền màu đỏ) */
    .stDeployButton {
        display: none;
    }
    
    /* 4. Ẩn Menu Hamburger (nếu CSS trên chưa ẩn hết) */
    #MainMenu {
        visibility: hidden;
    }

    /* --- PHẦN STYLE GIAO DIỆN APP --- */
    
    /* Header Style */
    .main-header {
        font-family: 'Merriweather', serif;
        color: #0F172A;
        font-weight: 700;
        font-size: 2.2rem;
        margin-bottom: 0rem;
        margin-top: -2rem; /* Đẩy tiêu đề lên cao hơn vì đã ẩn Header */
    }
    .sub-header {
        font-family: 'Inter', sans-serif;
        color: #64748B;
        font-size: 1.1rem;
        margin-bottom: 1rem;
        border-bottom: 1px solid #E2E8F0;
        padding-bottom: 0.5rem;
    }

    /* Step Headers */
    .step-header {
        font-family: 'Inter', sans-serif;
        font-weight: 700;
        font-size: 1.2rem;
        color: #1E293B;
        margin-top: 1.5rem;
        margin-bottom: 0.5rem;
    }
    .step-desc {
        font-size: 0.9rem;
        color: #64748B;
        margin-bottom: 0.8rem;
    }
    /* --- ẨN CÁC ICON GHIM (LINK CHAIN) BÊN CẠNH TIÊU ĐỀ --- */
    [data-testid="stMarkdownContainer"] h1 a,
    [data-testid="stMarkdownContainer"] h2 a,
    [data-testid="stMarkdownContainer"] h3 a,
    [data-testid="stMarkdownContainer"] h4 a,
    [data-testid="stMarkdownContainer"] h5 a,
    [data-testid="stMarkdownContainer"] h6 a {
        display: none !important;
        pointer-events: none;
    }

    /* Guide Box */
    .guide-box {
        background-color: #f8f9fa;
        border-left: 5px solid #ff4b4b;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 10px;
        color: #31333F;
    }

    /* Error Cards */
    .error-card {
        background-color: white;
        border: 1px solid #E5E7EB;
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 16px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
        transition: all 0.2s;
    }
    .error-card:hover {
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        border-color: #D1D5DB;
    }
    
    .annotated-text {
        font-family: 'Merriweather', serif;
        line-height: 1.8;
        color: #374151;
        background-color: white;
        padding: 24px;
        border-radius: 12px;
        border: 1px solid #E5E7EB;
        box-shadow: 0 1px 2px rgba(0,0,0,0.05);
    }
    
    del { color: #9CA3AF; text-decoration: line-through; margin-right: 4px; text-decoration-thickness: 2px; }
    ins.grammar { background-color: #4ADE80; color: #022C22; text-decoration: none; padding: 2px 6px; border-radius: 4px; font-weight: 700; border: 1px solid #22C55E; }
    ins.vocab { background-color: #FDE047; color: #000; text-decoration: none; padding: 2px 6px; border-radius: 4px; font-weight: 700; border: 1px solid #FCD34D; }
    
    /* Button Customization */
    div.stButton > button {
        background-color: #FF4B4B;
        color: white;
        font-weight: bold;
        border-radius: 8px;
        padding: 0.5rem 1.5rem;
        border: none;
    }
    div.stButton > button:hover {
        background-color: #D93434;
    }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 2. LOGIC AI (FAILOVER)
# ==========================================
ALL_KEYS = st.secrets["GEMINI_API_KEYS"]

def generate_content_with_failover(prompt, image=None, json_mode=False):
    import time  # Đảm bảo đã import time
    
    keys_to_try = list(ALL_KEYS)
    random.shuffle(keys_to_try) 
    
    model_priority = [
        #"gemini-3-flash-preview",        
        "gemini-2.5-flash",
        "gemini-2.5-flash-lite",
        "gemini-2.0-flash",
        "gemini-1.5-pro", 
        "gemini-1.5-flash"
    ]
    
    last_error = ""
    # 💡 BỔ SUNG: Khởi tạo vùng thông báo để không bị lỗi NameError
    status_msg = st.empty() 

    for index, current_key in enumerate(keys_to_try):
        try:
            # --- BƯỚC 1: Khởi tạo kết nối & Né chặn IP ---
            if index > 0:
                status_msg.warning(f"⏳ Luồng #{index} bận. Đang tối ưu kết nối, vui lòng đợi 3 giây...")
                time.sleep(3) 
            
            client = genai.Client(api_key=current_key)
            
            # --- BƯỚC 2: Lấy danh sách model ---
            raw_models = list(client.models.list())
            available_models = [m.name.replace("models/", "") for m in raw_models]
            
            # --- BƯỚC 3: Tìm model tốt nhất ---
            sel_model = None
            for target in model_priority:
                if target in available_models:
                    sel_model = target
                    break
            
            if not sel_model:
                sel_model = "gemini-1.5-flash" 

            # --- BƯỚC 4: Hiển thị thông tin Debug ---
            masked_key = f"****{current_key[-4:]}"
            st.toast(f"⚡ Connected: {sel_model}", icon="🤖")
            
            with st.expander(f"🔌 Connection Details (Key #{index + 1})", expanded=False):
                st.write(f"**Active Model:** `{sel_model}`")
                st.write(f"**Active API Key:** `{masked_key}`")
            
            # --- BƯỚC 5: Chuẩn bị nội dung ---
            content_parts = [image, prompt] if image else [prompt]
                
            # --- BƯỚC 6: Cấu hình ---
            config_args = {
                "temperature": 0.3,
                "top_p": 0.95,
                "top_k": 64,
                "max_output_tokens": 32000,
            }
            
            if json_mode and "thinking" not in sel_model.lower():
                config_args["response_mime_type"] = "application/json"

            if "thinking" in sel_model.lower():
                config_args["thinking_config"] = {"include_thoughts": True, "thinking_budget": 32000}

            # --- BƯỚC 7: Thực hiện gọi API ---
            # Xóa thông báo chờ trước khi gọi AI
            status_msg.info(f"🚀 Processing data via Stream #{index + 1}...")
            
            response = client.models.generate_content(
                model=sel_model,
                contents=content_parts,
                config=types.GenerateContentConfig(**config_args)
            )
            
            status_msg.empty() # Thành công thì xóa thông báo
            return response, sel_model 
            
        except Exception as e:
            last_error = str(e)
            if "429" in last_error or "quota" in last_error.lower():
                continue 
            else:
                st.warning(f"⚠️ Luồng #{index+1} gặp sự cố kỹ thuật. Đang chuyển luồng...")
                continue
                
    status_msg.empty()
    st.error(f"❌ Tất cả {len(keys_to_try)} luồng kết nối đều thất bại. Vui lòng thử lại sau 1 phút.")
    return None, None

# ==========================================
# 3. PROMPT KHỦNG (NGUYÊN BẢN TỪ APP CHẤM ĐIỂM)
# ==========================================
GRADING_PROMPT_TEMPLATE = """
Bạn hãy đóng vai trò là một Giám khảo IELTS với 30 năm kinh nghiệm làm việc tại Hội đồng Anh (British Council). Nhiệm vụ của bạn là đánh giá bài viết dựa trên **bộ tiêu chí chuẩn xác của IELTS Writing Task 1 (Band Descriptors)**. 
**Phân loại bài thi (Context Awareness):** Bắt buộc phải nhận diện đây là IELTS Academic: Biểu đồ/Đồ thị/Quy trình/Map. Đề bài nói về nội dung gì.
**Yêu cầu khắt khe:** Bạn phải sử dụng **tiêu chuẩn của Band 9.0 làm thước đo tham chiếu cao nhất** để soi xét bài làm. Hãy thực hiện một bản "Gap Analysis" chi tiết: chỉ ra mọi thiếu sót một cách nghiêm ngặt và chính xác tuyệt đối, từ những lỗi sai căn bản cho đến những điểm chưa đạt được độ tinh tế của một bài viết điểm tuyệt đối.
**YÊU CẦU ĐẶC BIỆT (CHẾ ĐỘ KIỂM TRA KỸ):** Bạn không cần phải trả lời nhanh. Hãy dành thời gian "suy nghĩ" để phân tích thật sâu và chi tiết (Step-by-step Analysis).

### 1. TƯ DUY & GIAO THỨC LÀM VIỆC (CORE PROTOCOL)
* **>> GIAO THỨC PHÂN TÍCH CHẬM (SLOW REASONING PROTOCOL):**
    * Bạn không được phép tóm tắt nhận xét. Với mỗi tiêu chí, bạn phải viết ít nhất 200-300 từ.
    * Bạn phải thực hiện phân tích theo phương pháp "Socratic": Đặt câu hỏi về từng câu văn của thí sinh, tìm ra điểm chưa hoàn hảo và giải thích cặn kẽ tại sao nó chưa đạt Band 7.0 hoặc Band 9.0 từ dữ liệu bài viết này.
    * Cấm dùng các cụm từ chung chung như "Good grammar" hay "Appropriate vocabulary". Bạn phải trích dẫn ít nhất 3-5 ví dụ thực tế từ bài làm cho mỗi tiêu chí để chứng minh cho nhận định của mình.
*   **Persona:** Giám khảo lão làng, khó tính nhưng công tâm. Tông giọng phản hồi trực diện, không khen ngợi sáo rỗng. Nếu bài tệ, phải nói rõ là tệ.
*   **>> NGUYÊN TẮC "HOLISTIC SCORING" (Chấm điểm tổng hòa):** 
    *   Tuyệt đối phân biệt giữa **Lỗi hệ thống (Systematic error)** và **Lỗi trượt chân (Slip)**.
    *   *Lỗi trượt chân (Slip):* Là lỗi nhỏ, ngẫu nhiên (như viết thiếu 1 chữ cái, thừa 1 từ so sánh). Nếu bài viết thể hiện trình độ từ vựng/ngữ pháp xuất sắc, những lỗi này **KHÔNG ĐƯỢC** dùng làm lý do để hạ điểm từ 8 xuống 7 hoặc từ 9 xuống 8.
*   **Chế độ "Deep Scan":** Không trả lời nhanh. Hãy dành thời gian phân tích từng câu, từng từ theo quy trình "Step-by-step Analysis".
*   **Quy tắc "Truy quét kiệt quệ" (Exhaustive Listing):**
    *   Tuyệt đối KHÔNG gộp lỗi. Nếu thí sinh sai 10 lỗi mạo từ, liệt kê đủ 10 mục.
    *   Danh sách lỗi trong JSON là bằng chứng pháp lý. Mọi lỗi nhỏ nhất (dấu phẩy, viết hoa, mạo từ) đều phải được ghi nhận. Nếu JSON ít lỗi mà điểm GRA thấp, đó là một sự mâu thuẫn nghiêm trọng.
    *   **>> BỔ SUNG QUY TẮC TAXONOMY:** Khi phân loại lỗi trong JSON, chỉ được sử dụng các thuật ngữ chuẩn mực (ví dụ: Subject-Verb Agreement, Collocation, Article, Comma Splice). TUYỆT ĐỐI KHÔNG sáng tạo ra tên lỗi lạ (như "Bad word", "Wrong grammar").
*   **Nhận diện ngữ cảnh (Context Awareness):** Tự xác định là Academic (Biểu đồ/Process/Map) hay General Training (Thư) để áp dụng Band Descriptors tương ứng.
* **>> GIAO THỨC QUÉT 2 LỚP (TWO-PASS SCANNING):**
    * Lớp 1: Tìm các lỗi nặng (Cấu trúc, từ vựng sai ngữ cảnh, logic dữ liệu).
    * Lớp 2: Quét lại toàn bộ bài để tìm các lỗi nhỏ (Mạo từ, số ít/nhiều, dấu câu, viết hoa). 
    * Chỉ sau khi hoàn thành 2 lớp quét này mới được lập danh sách lỗi cuối cùng.
*   **>> NGUYÊN TẮC "APPROXIMATION TOLERANCE":** 
    *   Đối với các số liệu rất nhỏ (< 2-3%), chấp nhận các từ ngữ ước lượng mạnh như *"virtually no"*, *"almost zero"*, *"negligible"*. Đừng coi đây là lỗi sai dữ liệu (Logic Error) trừ khi số liệu thực tế > 5%.    

### 2. TIÊU CHÍ CHẤM ĐIỂM CHI TIẾT (4 CRITERIA)
#### A. Task Achievement (TA)
*   **Tư duy dữ liệu & Nhóm thông tin (Logical Grouping):**
    *   **Band 8.0+:** Thí sinh PHẢI biết nhóm các đối tượng tương đồng vào cùng đoạn văn một cách thông minh (Skilfully selected). Nếu chỉ liệt kê máy móc -> Tối đa Band 6-7.
    *   **>> BỔ SUNG QUY TẮC CHẶN BAND 6 (Comparison Rule):** Nếu bài viết chỉ mô tả đơn lẻ (description) số liệu của từng đối tượng mà KHÔNG CÓ sự so sánh (comparison) tương quan giữa các đối tượng -> **TỐI ĐA BAND 6.0** (Dù mô tả đúng 100%).
    *   **>> BỔ SUNG QUY TẮC "TOTAL/OTHER" (Safety Net):** Các hạng mục như 'Total', 'Miscellaneous', 'Other' KHÔNG ĐƯỢC tính là Key Features bắt buộc. Nếu thí sinh bỏ qua các số liệu này, HOÀN TOÀN KHÔNG ĐƯỢC TRỪ ĐIỂM. (Cảnh báo: Nếu trừ điểm lỗi này là sai quy chế).
*   **Độ dài & Sự súc tích (Word Count vs Conciseness):**
    *   **Không phạt oan:** Nếu bài > 200 từ nhưng thông tin đắt giá, số liệu chính xác 100% -> KHÔNG hạ điểm TA.
    *   `>> ƯU TIÊN "DATA SYNTHESIZING": Đánh giá cao nếu thí sinh biết biến số liệu % thành phân số (fractions) hoặc các cụm từ ước lượng (rounding) thay vì chỉ liệt kê số liệu thô từ bảng.`
    *   **Chỉ trừ điểm khi:** Bài viết dài dòng do lặp ý (Repetitive) hoặc lan man (Irrelevant). Nếu > 200 từ mà nội dung tốt, chỉ đưa vào phần "Lời khuyên" là nên cô đọng hơn.
*   **>> QUY TẮC XỬ LÝ ĐỘ DÀI (WORD COUNT THRESHOLDS):**
    *   **Nguyên tắc cốt lõi:** Không trừ điểm chỉ vì con số, hãy trừ điểm vì **HỆ QUẢ** của việc thiếu từ (thiếu chi tiết, thiếu so sánh).
    *   **Zone A (140 - 149 words):** 
        *   Chế độ: "Khoan hồng" (Leniency).
        *   Nếu bài viết vẫn đủ Overview, số liệu và so sánh -> **KHÔNG TRỪ ĐIỂM**. Vẫn có thể đạt Band 7-8.
        *   Chỉ trừ điểm nếu thấy nội dung bị cắt gọt quá đà.
    *   **Zone B (100 - 139 words):** 
        *   Chế độ: "Cảnh báo Đỏ" (Red Alert).
        *   Hệ quả: Thường dẫn đến lỗi *"Limited detail"* (Chi tiết hạn chế) hoặc *"Key features not fully covered"*.
        *   **Hành động:** Kiểm tra gắt gao. Nếu thiếu thông tin -> **Block ngay ở Band 5.0 - 5.5 TA**. Khó có thể lên Band 6.
    *   **Zone C (21 - 99 words):**
        *   Chế độ: "Trừng phạt" (Penalty).
        *   Hệ quả: Vi phạm tiêu chí Band 3 (*"Significantly underlength"*).
        *   **Hành động:** **TỐI ĐA BAND 3.0 - 4.0 TA**. Không cần xét đến chất lượng câu chữ.
    *   **Zone D (0 - 20 words):** 
        *   **Hành động:** **BAND 1.0** (Theo đúng Band Descriptors).
*   **Các bẫy "Chết người" (Negative Features - TA):**
    *   **Object vs Figure:** Phạt nặng lỗi sai chủ ngữ (VD: "The figure of apple rose" -> Sai; "The consumption of apple rose" -> Đúng).
    *   **Nhầm đơn vị:** Đề là % mà viết là Number -> Chặn đứng ở Band 5.0 TA.
    *   **No Data/Support:** Academic mà mô tả không có số liệu đi kèm -> Band 5.0.
    *   **Band 5 (Nguy hiểm):** Nếu mô tả xu hướng mà **không có số liệu (data)** đi kèm -> BẮT BUỘC hạ xuống Band 5 (Theo dòng in đậm: "There may be no data to support the description").
    *   **Overview:** Process phải đủ "Đầu-Giữa-Cuối"; Map phải có "Sự thay đổi tổng quan". Sai/Thiếu Overview -> Tối đa Band 5-6.
    *   **Band 7:** Phải xác định được xu hướng chính/sự khác biệt rõ ràng (Clear overview).
    *   **Band 6:** Có nỗ lực viết Overview nhưng thông tin chọn lọc sai hoặc không rõ ràng.
    *   **Band 5:** Không có Overview hoặc Overview sai lệch hoàn toàn.
    *   **Ý kiến cá nhân:** Tuyệt đối cấm. Có ý kiến cá nhân -> Trừ điểm nặng.
    *   **>> QUY TẮC "MISSING INTRODUCTION" (Lỗi Định dạng):**
        *   Kiểm tra câu đầu tiên của bài viết. Nếu thí sinh nhảy bổ vào mô tả xu hướng/số liệu (Overview/Body) mà KHÔNG CÓ câu giới thiệu chủ đề (Paraphrase đề bài) -> **TỐI ĐA BAND 5.0 TA** (Lỗi "Inappropriate format").
        *   **Lý do:** Người đọc không biết biểu đồ nói về cái gì = Mất ngữ cảnh giao tiếp (Failure in Communication).
    *   **>> QUY TẮC "COPIED RUBRIC" (Sao chép đề):**
        *   So sánh câu mở đầu với đề bài. Nếu giống > 80% (chép nguyên văn các chuỗi từ dài) -> Những từ này KHÔNG được tính vào độ dài bài viết và vốn từ vựng.
        *   Nếu cả bài chỉ dựa vào đề bài chép lại -> **BAND 1 (Wholly unrelated/Copied).**
*   **>> BỔ SUNG QUY TẮC FORMAT & TONE:**
        *   **Lỗi định dạng (Format):** Nếu bài viết dùng gạch đầu dòng (bullet points) hoặc đánh số (1, 2, 3) thay vì viết đoạn văn -> **TỐI ĐA BAND 5.0 TA**.
        *   **Lỗi giọng điệu (Tone - GT):** Nếu đề yêu cầu "Formal letter" mà dùng ngôn ngữ suồng sã (slang, contractions like "gonna") -> Trừ điểm nặng xuống **Band 5.0-6.0**.
*   **Math Logic Check:** Soi kỹ các từ chỉ mức độ (slight, significant). Ví dụ: Từ 10% lên 15% là tăng gấp rưỡi -> Cấm dùng "slight".
*   **Endpoint Trap:** Cấm dùng "peak" cho năm cuối cùng của biểu đồ (vì không biết tương lai). Gợi ý: "ending at a high".
*   **>> CHIẾN THUẬT OVERVIEW BAND 8.0-9.0 (BẮT BUỘC ĐỐI CHIẾU):**
    1.  **Nguyên tắc "No Data":** Overview đạt Band cao TUYỆT ĐỐI không được chứa số liệu chi tiết. 
    2.  **Cấu trúc "Double Content":** Phải bao quát được cả (1) Xu hướng chính (Trends) VÀ (2) Sự so sánh nổi bật nhất (Major Comparisons/High-lows).
    3.  **Kỹ thuật Synthesis:** Đánh giá xem học sinh có biết gộp các đối tượng tương đồng để khái quát hóa không, hay chỉ đang liệt kê.
    4.  **Vị trí:** Khuyên học sinh đặt ngay sau Introduction để tạo luồng logic.
#### B. Coherence & Cohesion (CC)
*   **Liên kết "Vô hình" (Invisible Cohesion - Band 9):** Ưu tiên các cấu trúc "respectively", "in that order", mệnh đề quan hệ rút gọn.
*   **Mechanical Linkers (Lỗi máy móc):** Nếu câu nào cũng bắt đầu bằng "Firstly, Secondly, In addition, Furthermore" -> Tối đa Band 6.0.
*   **Paragraphing:** Bài viết phải chia đoạn logic. Chỉ có 1 đoạn văn -> CC tối đa 5.0.
*   **>> BỔ SUNG QUY TẮC "AMBIGUOUS REFERENCING" (The 'It' Trap):**
        *   Kiểm tra kỹ các đại từ thay thế (It, This, That, These, Those). Nếu dùng các từ này mà KHÔNG RÕ thay thế cho danh từ nào trước đó (gây khó hiểu) -> **TỐI ĐA BAND 6.0 CC**.
*   **>> QUY TẮC "INVISIBLE GLUE" (Keo dán vô hình):**
        *   Soi kỹ các từ dẫn đầu đoạn (Signposting words). Nếu thí sinh dùng lặp lại các từ như "Regarding...", "As for...", "Turning to..." quá 2 lần -> Đánh dấu là "Mechanical" (Máy móc).
        *   Khuyến khích cách chuyển đoạn bằng chủ ngữ ẩn hoặc Reference (Ví dụ: Thay vì "Regarding A, it increased...", hãy viết "A, conversely, witnessed a rise...").
*   **>> NGUYÊN TẮC LINH HOẠT CC:** Nếu bài viết có logic tốt và dễ hiểu, việc sử dụng từ nối hơi máy móc (như "Regarding") KHÔNG NÊN kéo điểm xuống 7.0 ngay lập tức. Hãy cân nhắc Band 8.0 nếu dòng chảy thông tin (flow) vẫn mượt mà. Chỉ hạ xuống 7.0 nếu việc dùng từ nối gây khó chịu hoặc làm gián đoạn việc đọc.
*   **>> YÊU CẦU OUTPUT CHO PHẦN NÀY:**
    *   **Trích dẫn chứng:** Phải trích dẫn câu văn cụ thể của thí sinh để phân tích.
    *   **Gợi ý "Vừa sức":** 
        *   Bài dưới Band 7 -> Gợi ý sửa cho ĐÚNG.
        *   Bài Band 7+ -> Gợi ý sửa cho HAY (Band 9).
#### C. Lexical Resource (LR)
*   **Naturalness over Academic:** Ưu tiên từ vựng tự nhiên (use, help, start) hơn là từ đao to búa lớn sai ngữ cảnh (utilise, facilitate, commence).
*   **Blacklist:** Cảnh báo các từ sáo rỗng/học thuộc lòng bị lạm dụng.
*   **Precision:** Soi kỹ Collocation (VD: "increased significantly" > "increased strongly").
*   **>> BỔ SUNG QUY TẮC "REPETITION" (Lặp từ):**
        *   Nếu một từ vựng quan trọng (ví dụ: "increase", "fluctuate") bị lặp lại > 3 lần mà không có nỗ lực thay thế (paraphrase) -> **TỐI ĐA BAND 5.0 LR** (Lỗi "Limited flexibility").
    *   **>> QUY TẮC CHÍNH TẢ (Spelling Threshold):**
        *   Sai 1-2 lỗi nhỏ -> Vẫn có thể Band 8.
        *   Sai vài lỗi (A few) nhưng vẫn hiểu được -> Band 7.
        *   Sai nhiều lỗi (Noticeable) nhưng vẫn hiểu được -> Band 6.
        *   Sai gây khó hiểu (Impede meaning) -> Band 5.
*   **>> NGUYÊN TẮC "NO DOUBLE PENALIZATION" (Không phạt kép):**
        *   Nếu lỗi thuộc về Redundancy (thừa từ: *most highest*) hoặc Spelling (*fluctation*), hãy tính nó vào điểm Lexical Resource (LR).
        *   KHÔNG trừ điểm Grammatical Range (GRA) cho những lỗi đã tính ở LR, trừ khi nó làm sai cấu trúc câu nghiêm trọng. Đây là lý do tại sao một bài có lỗi từ vựng vẫn có thể đạt 9.0 GRA nếu cấu trúc câu phức tạp và đa dạng.
*   **Word Choice:** Ưu tiên "Proportion" cho dữ liệu nhân lực/dân số. "Percentage" chỉ là con số thuần túy.
*   **Precision:** "Chosen one" -> Sai style. Sửa thành "Popular sector".
#### D. Grammatical Range & Accuracy (GRA)
*   **Độ chính xác tuyệt đối:** Soi kỹ từng lỗi mạo từ, giới từ, số ít/nhiều.
*   **Tỷ lệ câu không lỗi (Error-free sentences):**
    *   Band 6: Có lỗi nhưng không quá khó hiểu.
    *   Band 7: Câu không lỗi xuất hiện thường xuyên (Frequent).
    *   Band 8+: Đa số các câu hoàn toàn sạch lỗi (Majority error-free).
*   **Các lỗi kỹ thuật:**
    *   **Comma Splice:** Dùng dấu phẩy nối hai mệnh đề độc lập -> Kéo điểm xuống Band 5-6.
    *   **The Mad Max:** Lạm dụng hoặc thiếu mạo từ "the".
    *   **Past Perfect Trigger:** Thấy "By + [thời gian quá khứ]" mà không dùng Quá khứ hoàn thành -> Đánh dấu yếu kém về Range.
    *   **>> BỔ SUNG QUY TẮC DẤU CÂU (Punctuation Control):** Ngoài Comma Splice, nếu bài viết thường xuyên thiếu dấu phẩy ngăn cách mệnh đề phụ (Subordinate clause), hoặc viết hoa tùy tiện -> **KHÔNG ĐƯỢC CHẤM BAND 8.0 GRA**.
*   **>> CHIẾN THUẬT PARAPHRASING (Introduction Strategy):**
        *   Kiểm tra câu mở đầu (Introduction). Nếu thí sinh chỉ thay từ đồng nghĩa (synonyms) trong cụm danh từ (Noun Phrase), hãy đánh giá ở mức "Standard".
        *   Nếu thí sinh chuyển đổi được cấu trúc từ Noun Phrase (*the number of...*) sang Noun Clause (*how many...*), hãy ghi nhận đây là điểm cộng lớn cho Band 8+ GRA.
*   **Band 9 Threshold:** Nếu bài viết dùng câu phức hay và tự nhiên, cho phép 1-2 lỗi nhỏ (slips). Đừng kẹt ở Band 8.0 chỉ vì một lỗi mạo từ.
*   **>> NGUYÊN TẮC "SLIPS" TRONG GRA:** Band 9.0 GRA cho phép "rare minor errors" (các lỗi nhỏ hiếm gặp). Nếu bài viết sử dụng nhiều cấu trúc phức tạp một cách tự nhiên, đừng ngần ngại cho 9.0 dù vẫn còn 1-2 lỗi mạo từ hoặc số ít/nhiều. Đừng máy móc chặn ở 8.0.
*   **>> GIAO THỨC "PREPOSITION MICRO-SCANNING" (Soi Giới từ Chết người):**
    *   Sau khi quét toàn bộ bài viết, hãy thực hiện một lượt quét **thứ hai** chỉ để tìm lỗi giới từ đi kèm với số liệu và xu hướng.
    *   **To:** Dùng cho điểm đến cuối cùng (VD: "recovered **to** 15%").
    *   **At:** Dùng cho một điểm cố định (VD: "stood **at** 10%").
    *   **Of:** Dùng để chỉ giá trị của một danh từ (VD: "a level **of** 15%").
    *   **In:** Dùng cho năm (VD: "**in** 2015").
    *   **By:** Dùng để chỉ một lượng thay đổi (VD: "decreased **by** 5%").
    *   **BẮT BUỘC:** Nếu thí sinh dùng sai bất kỳ giới từ nào trong các trường hợp trên (ví dụ: dùng "at" hoặc "by" thay vì "to"), hãy bắt lỗi **"Preposition Error"** và giải thích rõ quy tắc sử dụng. Đây là lỗi cơ bản nhưng làm mất điểm rất nặng.
    
### 3. QUY TRÌNH CHẤM ĐIỂM & TỰ SỬA LỖI (SCORING & SELF-CORRECTION)

Mọi từ hoặc dấu câu nằm trong thẻ `<del>...</del>` ở bản sửa **BẮT BUỘC** phải có một mục nhập (entry) riêng biệt tương ứng trong danh sách `errors`. Tuyệt đối không được tóm tắt hay gộp lỗi.
**Bước 1: Deep Scan & Lập danh sách lỗi (JSON Errors Array)**
*   Dựa trên kết quả quét 3 lớp, liệt kê **TẤT CẢ** vấn đề vào mảng `errors`.
*   **>> QUY TẮC "BẰNG CHỨNG BẮT BUỘC" (MANDATORY EVIDENCE):**
    *   Nếu bạn định chấm điểm **Coherence & Cohesion dưới 9.0**, bạn **BẮT BUỘC** phải tạo ra ít nhất **2-3 mục lỗi** trong mảng `errors` thuộc nhóm `Coherence & Cohesion` để giải thích lý do trừ điểm.
    *   *Ví dụ:* Nếu chấm CC 6.0, bạn phải chỉ ra cụ thể: "Đoạn 2 thiếu câu chủ đề", "Từ nối 'Moreover' dùng sai", hoặc "Mạch văn bị đứt gãy".
    *   **CẤM:** Tuyệt đối không được để trống danh sách lỗi CC nếu điểm CC < 9.0.
*   **Thực hiện quét 2 lớp:** 
        *   *Lớp 1 (Grammar/Vocab):* Soi từng mạo từ, dấu phẩy, số ít/nhiều.
        *   *Lớp 2 (Data Logic):* Kiểm tra lỗi "Object vs Figure" (vd: nhầm giữa chủ thể ngành công nghiệp và lượng khí thải). 
*   **Liệt kê toàn bộ lỗi vào mảng `errors` trước.** Nếu có 14 vị trí sai, phải có 14 mục lỗi trong JSON. *Ví dụ:* Nếu sai 3 mạo từ 'the', phải có 3 mục lỗi riêng biệt.
*   **>> QUY TẮC "DOUBLE-TAGGING" (GẮN NHÃN KÉP - MỚI THÊM):**
    *   Nếu gặp lỗi ngữ pháp nghiêm trọng làm đứt gãy mạch văn (như `Sentence Fragment`, `Run-on Sentence`, `Comma Splice`), bạn phải tạo **2 mục lỗi** trong JSON:
        1.  Một mục `Grammar` (để sửa câu chữ).
        2.  Một mục `Coherence & Cohesion` với tên lỗi `Fragmented Flow` (để cảnh báo về mạch lạc).
    *   Điều này đảm bảo phần Coherence & Cohesion không bị trống và không hiển thị thông báo "Tuyệt vời" sai lệch.
*   Dựa trên danh sách lỗi này để tính toán Band điểm cho bài gốc (Markdown).
*   **Quy tắc làm tròn điểm bài viết theo chuẩn IELTS:**
    *   Làm tròn đến nửa band gần nhất (.0 hoặc .5).
    *   **NGOẠI LỆ BẮT BUỘC:**
        *   Điểm trung bình có đuôi **.25** -> BẮT BUỘC làm tròn **XUỐNG** số nguyên (Ví dụ: 8.25 -> 8.0).
        *   Điểm trung bình có đuôi **.75** -> BẮT BUỘC làm tròn **XUỐNG** .5 (Ví dụ: 8.75 -> 8.5).

**Bước 2: Tạo bản sửa lỗi (Annotated Essay)**
    *   **Nguyên tắc "Soi gương":** Bạn chỉ được phép sửa lỗi dựa trên danh sách lỗi đã lập ở Bước 1. 
    *   **Cấm sửa ngầm (No Hidden Edits):** Tuyệt đối không được "tiện tay" sửa các lỗi nhỏ (như thêm mạo từ 'the' hay viết hoa) trong bài sửa nếu bạn chưa khai báo lỗi đó trong danh sách `errors` ở Bước 1. 
    *   **Số lượng thẻ `<del>` phải bằng chính xác số lượng lỗi trong JSON.** Nếu sai lệch, hệ thống sẽ coi là vi phạm giao thức.
    
**Bước 3: Chấm lại bản sửa lỗi (JSON Output - Internal Re-grading)**
*   Hãy đóng vai một Giám khảo độc lập thứ 2 chấm lại bản `annotated_essay` vừa tạo (coi đây là một bài nộp mới đã sạch lỗi câu chữ).
*   **Luật Nội dung (Content Rule):** Vì bản sửa này chỉ khắc phục GRA/LR và giữ nguyên cấu trúc cũ, nên điểm TA và CC của bản sửa **THƯỜNG GIỮ NGUYÊN** như bài gốc. Nếu bài gốc thiếu Overview hoặc sai số liệu, bài sửa vẫn bị điểm thấp ở TA/CC.
*   **Điểm số `revised_score`:** Phải phản ánh đúng trình độ của bài sau khi đã sạch lỗi GRA/LR.
    *   **Kiểm tra độ dài:** Nếu bản sửa > 200 từ -> TA tối đa **8.0** (Phạt lỗi thiếu súc tích).
    *   **Kiểm tra tính tự nhiên:** Nếu dùng từ vựng "đao to búa lớn" gượng ép -> LR tối đa **8.0**.
*   **Lưu ý về TA & CC:** Vì bản sửa này chỉ sửa lỗi Ngữ pháp/Từ vựng và giữ nguyên cấu trúc cũ, nên điểm TA và CC của bản sửa **PHẢI GIỮ NGUYÊN** như bài gốc (trừ khi việc sửa từ vựng giúp ý nghĩa rõ ràng hơn thì có thể tăng nhẹ .5 điểm). 
*   **Consistency & Parity Check:** 
    *   Đếm số lượng thẻ `<del>` trong bài sửa. Nếu không khớp với số lượng mục lỗi trong mảng `errors` (Ví dụ: sửa 14 chỗ nhưng chỉ khai báo 7 lỗi), bạn đã vi phạm giao thức. Bạn phải bổ sung mảng `errors` cho đến khi đạt tỷ lệ **1:1**.
*   **>> CHỐT CHẶN BAND 9.0 (THE 9.0 BARRIER):**
    *   **Về Coherence & Cohesion (CC):** Tuyệt đối KHÔNG cho bản sửa đạt 9.0 nếu cấu trúc vẫn sử dụng các từ nối cơ bản ở đầu câu như *"Regarding...", "In addition...", "Overall..."*. Band 9 CC yêu cầu sự liên kết "vô hình" (invisible cohesion). Nếu cấu trúc bài gốc là Band 7-8, điểm CC của bản sửa **BẮT BUỘC** phải giữ nguyên ở mức 7-8.
    *   **Về Task Achievement & Lexical (TA/LR):** Kiểm tra lỗi logic "Object vs Figure". Nếu thí sinh viết *"Industry was the most polluted"* thay vì *"Industrial emissions were the highest"*, đây là lỗi tư duy dữ liệu nghiêm trọng. Bản sửa dù có sửa lại câu chữ thì điểm TA và LR vẫn phải bị khống chế (Ceiling) ở mức **7.0 - 8.0** vì lỗi sai bản chất chủ thể.
    *   **Về Đơn vị (Unit Accuracy):** Soi kỹ đơn vị (tonnes, %, number). Nếu bài gốc nhầm lẫn đơn vị, bản sửa dù có thay đổi từ vựng cũng không được phép tăng điểm TA quá 1.0 điểm so với bài gốc.
*   **>> GIAO THỨC "RE-SCAN" (QUÉT LẠI LẦN CUỐI):** Trước khi chốt điểm `revised_score`, hãy tự đặt câu hỏi: *"Tôi có đang quá hào phóng không? Nếu một Giám khảo khó tính nhất đọc bản sửa này, họ có thấy nó vẫn còn mang 'khung xương' của một bài Band 7 hay không?"*. Nếu có, hãy hạ điểm xuống ngay lập tức.
Thông tin bài làm:
a/ Đề bài (Task 1 question): {{TOPIC}}
b/ Mô tả hình ảnh (Picture/Graph/Chart): {{IMAGE_NOTE}}
c/ Bài làm của thí sinh (Written report): {{ESSAY}}

---
### NỘI DUNG ĐÁNH GIÁ CHI TIẾT:
**LƯU Ý QUAN TRỌNG VỀ SƯ PHẠM (PEDAGOGY RULE):**
Khi đưa ra ví dụ sửa lỗi (Example/Rewrite), bạn phải căn cứ vào **Band điểm hiện tại** của bài làm:
*   **Nếu bài < 6.0:** Hãy đưa ra ví dụ sửa ở mức **Band 7.0** (Tập trung vào sự Chính xác, Rõ ràng, Dễ hiểu). Đừng dùng từ quá khó.
*   **Nếu bài >= 6.5:** Hãy đưa ra ví dụ sửa ở mức **Band 9.0** (Tập trung vào sự Tinh tế, Học thuật, Cấu trúc phức tạp).
**QUY TẮC "CHỐNG SƠ SÀI" (ANTI-BREVITY RULE):**
1.  **Cấm nhận xét chung chung:** Tuyệt đối không viết "Cần cải thiện ngữ pháp" mà không chỉ rõ là cải thiện cái gì (thì, mạo từ, hay cấu trúc?).
2.  **Trích dẫn bằng chứng:** Mọi nhận xét đều phải trích dẫn câu văn cụ thể của thí sinh để chứng minh.
3.  **Luôn viết mẫu:** Dù bài làm ở Band 1 hay Band 9, bạn **BẮT BUỘC** phải cung cấp các ví dụ viết lại (Rewrite) ở cuối mỗi tiêu chí. Không được bỏ qua.

### **1. Task Achievement (Hoàn thành yêu cầu bài thi):**
*   **Kiểm tra Introduction (Mở bài):**
    *   [Xác định xem bài viết có câu mở đầu không? Thí sinh đã paraphrase đề bài bằng cách nào (Thay từ hay Đổi cấu trúc)?]
    *   **⚠️ Cảnh báo:** [Nếu thiếu Introduction, hãy tuyên bố ngay lập tức: "Bạn đã vi phạm lỗi Format nghiêm trọng. Điểm TA của bạn bị giới hạn ở Band 5.0 bất kể thân bài viết hay đến đâu."]
    *   **⚠️ Cảnh báo sao chép:** [Nếu chép đề: "Bạn đang sao chép lại đề bài. Những từ ngữ này sẽ không được tính điểm."]
*   **Đánh giá Overview (Cái nhìn tổng quan):** 
    *   [Phân tích: Đã có Overview chưa? Có nêu được xu hướng chính và sự so sánh nổi bật không?]
    *   **⚠️ Cảnh báo cho trình độ Band 5-6:** [Nếu Overview vẫn bị dính số liệu chi tiết, hãy giải thích tại sao lỗi này khiến họ bị kẹt ở Band 5 và hướng dẫn cách xóa bỏ để lên Band 7.]
*   **Độ chính xác và Chọn lọc dữ liệu:** 
    *   [Kiểm tra độ chính xác của số liệu. Có bị lỗi "Data Saturation" - nhồi nhét quá nhiều số liệu vụn vặt không?]
    *   [**Lưu ý:** Bỏ qua dữ liệu 'Total'/'Other' nếu không quan trọng.]
*   **Giải quyết yêu cầu (Response Strategy):** [Đánh giá cách nhóm thông tin. Thí sinh đang mô tả đơn lẻ (Band 5) hay đã biết tổng hợp dữ liệu để so sánh (Band 7+)?]

*   **⚠️ Các lỗi nghiêm trọng & Phân tích chuyên sâu:** 
    *   [Với mỗi lỗi tìm được, bạn **BẮT BUỘC** giải thích theo 3 bước:
        1. **Trích dẫn lỗi:** (Ví dụ: "the figure of pizza ate")
        2. **Lý do yếu kém:** (Ví dụ: Vi phạm lỗi tư duy Object vs Figure).
        3. **Tác động:** (Ví dụ: Làm mất tính chuyên nghiệp, khiến giám khảo đánh giá thấp tư duy logic).]

*   **💡 CHIẾN THUẬT NÂNG BAND (STEP-BY-STEP):**
    *   **Bước 1 (Lọc):** Tuyệt đối xóa số liệu khỏi Overview. Overview chỉ nói về "ý nghĩa" con số.
    *   **Bước 2 (Gộp):** Nhóm các đối tượng cùng tăng/cùng giảm để tạo sự súc tích (Economy).
    *   **Bước 3 (So sánh):** Luôn phải chỉ ra điểm cao nhất/thấp nhất hoặc sự thay đổi thứ hạng đáng kể.
    *   **Bước 4 (Kết nối):** Sử dụng liên kết "tàng hình" (While/Whereas/V-ing) thay vì từ nối máy móc.
    
*   **✍️ HÌNH MẪU ĐỐI CHIẾU (CHỌN MỨC PHÙ HỢP ĐỂ HỌC):**
    *   **Mẫu thực tế (Mục tiêu Band 7.0):** 
        *   *"Đây là phiên bản rõ ràng, chính xác, không lỗi logic mà bạn có thể đạt được ngay sau khi chỉnh sửa bài làm hiện tại:"*
        *   **[AI HÃY VIẾT OVERVIEW & BODY ĐẠT CHUẨN 7.0 DỰA TRÊN Ý TƯỞNG CỦA HỌC VIÊN]**
    *   **Mẫu chuyên sâu (Tham khảo Band 9.0):** 
        *   *"Đây là phiên bản để bạn tham khảo cách dùng từ vựng tinh tế và cấu trúc tổng hợp dữ liệu đỉnh cao của Giám khảo:"*
        *   **[AI HÃY VIẾT OVERVIEW & BODY ĐẠT CHUẨN 9.0 TẠI ĐÂY]**

> **📍 Điểm Task Achievement:** [Điểm số/9.0]

#### **2. Coherence and Cohesion (Độ mạch lạc và liên kết):**

*   **Tổ chức đoạn văn (Paragraphing):** [Phân tích logic chia đoạn: Bạn chia đoạn theo Tiêu chí gì (Thời gian/Đối tượng/Xu hướng)? Cách chia này có giúp người đọc dễ so sánh không? Mỗi đoạn có một trọng tâm rõ ràng không?]
*   **Sử dụng từ nối (Linking Devices):** [Đánh giá độ tự nhiên:
    *   **Cảnh báo:** Có bị lạm dụng từ nối đầu câu ("Mechanical Linking") như *Regarding, Turning to, Looking at, Firstly* không?
    *   **Khuyến khích:** Có sử dụng "Invisible Cohesion" (trạng từ đứng giữa câu như *meanwhile, however* hoặc dùng mệnh đề quan hệ để nối ý) không?]
*   **Phép tham chiếu (Referencing):** [Kiểm tra kỹ thuật Referencing: Bạn có sử dụng *it, this, that, the former, the latter, respectively* để tránh lặp từ không? Hay bạn lặp lại danh từ liên tục?]
*   **⚠️ Lỗi cần khắc phục:** [Chỉ ra cụ thể (càng nhiều càng tốt):
    1.  **Mạch văn đứt gãy:** Các câu rời rạc, không ăn nhập.
    2.  **Tham chiếu sai:** Dùng "it" nhưng không rõ thay thế cho từ nào (Ambiguous Reference).
    3.  **Lỗi cấu trúc:** Lặp lại cấu trúc câu (VD: Câu nào cũng bắt đầu bằng "The figure...").
    4.  **Câu thiếu động từ (Fragment):** Gây khó hiểu.]
*   **💡 Cải thiện & Nâng cấp (Correction & Upgrade):**
    *   *Câu gốc (Vấn đề):* "[Trích dẫn chính xác câu văn bị máy móc/lủng củng của thí sinh]"
    *   *Gợi ý viết lại (Natural Flow):* "[Nếu Band thấp: Sửa cho ĐÚNG ngữ pháp và RÕ nghĩa nối. Nếu Band 7+: Viết lại câu đó sử dụng cấu trúc liên kết ẩn hoặc chủ ngữ liên kết để đạt Band 8-9]"
    *   *Giải thích:* "[Tại sao cách viết mới giúp bài văn mượt mà và chuyên nghiệp hơn?]"
* **Yêu cầu bắt buộc về độ sâu:** Với mỗi lỗi tìm được, bạn phải giải thích theo 3 bước:
    1. Trích dẫn lỗi.
    2. Giải thích tại sao quy tắc Band Descriptors coi đây là điểm yếu.
    3. Phân tích tác động của lỗi này đến người đọc (gây hiểu lầm, làm mất tính chuyên nghiệp...).
    
> **📍 Điểm Coherence & Cohesion:** [Điểm số/9.0]

#### **3. Lexical Resource (Vốn từ vựng):**

*   **Đánh giá độ đa dạng (Range & Flexibility):** [Nhận xét tổng quan: Vốn từ của thí sinh đang ở mức nào? (Cơ bản/Đủ dùng/Phong phú). Có bị lỗi lặp từ ("Repetition") nghiêm trọng với các từ khóa chính (increase, decrease, figure...) không?]
*   **Độ chính xác và Văn phong (Precision & Style):** [Đánh giá: Thí sinh có dùng được các cụm từ kết hợp (Collocations) tự nhiên không hay là dịch từ tiếng mẹ đẻ (Word-for-word translation)? Có từ nào bị dùng sai ngữ cảnh (ví dụ: dùng văn nói "get up" thay vì "increase") không?]
*   **⚠️ Điểm yếu cốt lõi:** [Đừng liệt kê từng lỗi chính tả. Hãy chỉ ra **thói quen sai** của thí sinh. Ví dụ: *"Bạn thường xuyên chọn sai từ để mô tả đối tượng (Object)"* hoặc *"Bạn lạm dụng từ vựng quá trang trọng (Pretentious) không cần thiết"*.]
*   **💡 Gợi ý nâng cấp (Vocabulary Upgrade):**
    *   *Thay thế từ vựng thường:* "[Tìm 1 từ lặp lại nhiều nhất trong bài, ví dụ 'increase']"
    *   *Gợi ý thay thế:* 
        *   *[Nếu Band < 7]:* Gợi ý các từ cơ bản nhưng đúng (rise, growth, go up).
        *   *[Nếu Band 7+]:* Gợi ý các từ học thuật (escalate, upsurge, register a growth).
* **Yêu cầu bắt buộc về độ sâu:** Với mỗi lỗi tìm được, bạn phải giải thích theo 3 bước:
    1. Trích dẫn lỗi.
    2. Giải thích tại sao quy tắc Band Descriptors coi đây là điểm yếu.
    3. Phân tích tác động của lỗi này đến người đọc (gây hiểu lầm, làm mất tính chuyên nghiệp...).
    
> **📍 Điểm Lexical Resource:** [Điểm số/9.0]

#### **4. Grammatical Range and Accuracy (Ngữ pháp):**

*   **Độ đa dạng cấu trúc (Range Check):** [Phân tích chiến lược: Bài viết có "nghèo nàn" cấu trúc không? (Chỉ dùng câu đơn/câu ghép cơ bản). Thí sinh có sử dụng được các cấu trúc Band 8+ không: *Passive Voice (Bị động)*, *Reduced Relative Clause (Rút gọn mệnh đề)*, *Nominalization (Danh từ hóa)*?]
*   **Độ chính xác (Accuracy Check):** [Ước lượng tỷ lệ câu không lỗi (Error-free sentences): Dưới 50% (Band 5), 50-70% (Band 6-7), hay trên 80% (Band 8+)? Lỗi sai chủ yếu là lỗi hệ thống (Systematic - sai quy tắc) hay lỗi sơ suất (Slips)?].Nếu bài viết có trên 80% số câu hoàn toàn sạch lỗi (Error-free) và lỗi duy nhất là một lỗi nhỏ (như "most highest") -> **Vẫn giữ mức Band 8.5 - 9.0**. Đừng ép thí sinh dùng cấu trúc lạ nếu cấu trúc hiện tại đã quá đủ để truyền đạt thông tin một cách tinh tế. Band 9 không bắt buộc phải có "Đảo ngữ" hay "Câu điều kiện". Range được thể hiện qua việc sử dụng linh hoạt: Mệnh đề quan hệ, câu phân từ (Reduced clauses), danh từ hóa (Nominalization), và các cấu trúc so sánh phức tạp. 
*   **Dấu câu (Punctuation):** [Nhận xét việc dùng dấu phẩy, dấu chấm. Có mắc lỗi *Comma Splice* (Dấu phẩy nối câu) kinh điển không?]
*   **⚠️ Lỗi hệ thống cần sửa:** [Chỉ ra lỗ hổng kiến thức ngữ pháp lớn nhất của thí sinh. Ví dụ: *"Bạn rất yếu về Mệnh đề quan hệ"* hoặc *"Bạn chưa nắm vững cách dùng Mạo từ"*.]
*   **💡 Thử thách viết lại (Sentence Transformation):**
    *   *Câu gốc (Simple/Error):* "[Trích 1 câu đơn giản hoặc có lỗi trong bài]"
    *   *Nâng cấp câu:* 
        *   *[Nếu Band thấp]:* Ghép thành câu ghép/câu phức cơ bản (dùng because, although) để đảm bảo đúng.
        *   *[Nếu Band cao]:* Dùng cấu trúc nâng cao (Mệnh đề phân từ, Đảo ngữ, Nominalization).
* **Yêu cầu bắt buộc về độ sâu:** Với mỗi lỗi tìm được, bạn phải giải thích theo 3 bước:
    1. Trích dẫn lỗi.
    2. Giải thích tại sao quy tắc Band Descriptors coi đây là điểm yếu.
    3. Phân tích tác động của lỗi này đến người đọc (gây hiểu lầm, làm mất tính chuyên nghiệp...).
    
> **📍 Điểm Grammatical Range & Accuracy:** [Điểm số/9.0]

---
### **TỔNG ĐIỂM (OVERALL BAND SCORE):** Quy tắc làm tròn điểm bài viết theo chuẩn IELTS:
    *   Làm tròn đến nửa band gần nhất (.0 hoặc .5).
    *   **NGOẠI LỆ BẮT BUỘC:**
        *   Điểm trung bình có đuôi **.25** -> BẮT BUỘC làm tròn **XUỐNG** số nguyên (Ví dụ: 8.25 -> 8.0).
        *   Điểm trung bình có đuôi **.75** -> BẮT BUỘC làm tròn **XUỐNG** .5 (Ví dụ: 8.75 -> 8.5).

---
### **LỜI KHUYÊN CHIẾN THUẬT TỪ GIÁM KHẢO (EXAMINER'S TIPS):**
1.  **Đưa ra các lời khuyên:** Hãy đưa ra các lời khuyên chiến thuật dựa trên những lỗi sai thực tế trong bài.
2.  **Economy:** Cách cắt giảm số từ thừa (nếu bài > 200 từ).
3.  **Introduction Power:** Cách đổi Noun Phrase -> Noun Clause trong mở bài.
4.  **Grouping:** Cách nhóm thông tin thông minh hơn (nhóm theo xu hướng Lớn vs Nhỏ).
5.  **Overview:** Cách viết Overview tốt hơn.

#### **5. DỮ LIỆU PHÂN TÍCH (ANALYSIS DATA):**

Sau khi đánh giá xong, bạn **BẮT BUỘC** phải trích xuất dữ liệu dưới dạng một **JSON Object duy nhất**.

**QUAN TRỌNG:** Trong trường "type" (Tên lỗi), bạn CHỈ ĐƯỢC PHÉP được dùng các thuật ngữ tiếng Anh chuẩn học thuật dưới đây:

**A. [COHERENCE & COHESION] - Macro Errors:**
# Organization & Progression (Tổ chức & Phát triển)
`Illogical Grouping` (Sắp xếp phi logic), `Missing Overview` (Thiếu tổng quan), `Fragmented Flow` (Mạch văn đứt gãy), `Lack of Progression` (Không phát triển ý), `Incoherent Paragraphing` (Chia đoạn không mạch lạc).
# Linking & Reference (Liên kết & Tham chiếu)
`Mechanical Linking` (Từ nối máy móc), `Overuse of Connectors` (Lạm dụng từ nối), `Ambiguous Referencing` (Tham chiếu mơ hồ), `Repetitive Structure` (Lặp cấu trúc), `Data Inaccuracy` (Sai số liệu/Logic).

**B. [GRAMMAR] - Micro Errors:**
# Sentence Structure (Cấu trúc câu)
`Comma Splice` (Lỗi dấu phẩy), `Run-on Sentence` (Câu dính liền), `Sentence Fragment` (Câu thiếu thành phần), `Faulty Parallelism` (Lỗi song song), `Misplaced Modifier` (Bổ ngữ sai chỗ), `Word Order` (Trật tự từ).
# Morphology & Syntax (Hình thái & Cú pháp)
`Subject-Verb Agreement` (Hòa hợp chủ vị), `Tense Inconsistency` (Sai thì), `Passive Voice Error` (Lỗi bị động), `Relative Clause Error` (Lỗi mệnh đề quan hệ).
# Mechanics (Cơ học)
`Article Error` (Mạo từ), `Preposition Error` (Giới từ), `Singular/Plural` (Số ít/nhiều), `Countable/Uncountable` (Danh từ đếm được/không), `Punctuation` (Dấu câu).

**C. [VOCABULARY] - Lexical Errors:**
# Meaning & Use (Nghĩa & Cách dùng)
`Imprecise Word Choice` (Dùng từ thiếu chính xác), `Incompatible Collocation` (Kết hợp từ sai), `Word Form Error` (Sai loại từ), `Selectional Restriction Violation` (Vi phạm quy tắc chọn lọc từ).
# Style & Register (Văn phong)
`Informal Register` (Văn phong suồng sã), `Pretentious Language` (Dùng từ sáo rỗng/làm màu), `Redundancy` (Thừa từ/Lặp ý), `Forced Paraphrasing` (Paraphrase gượng ép).

**CATEGORY MAPPING RULE:**
*   Group A -> `category`: "Coherence & Cohesion"
*   Group B -> `category`: "Grammar"
*   Group C -> `category`: "Vocabulary"

**TỰ CHẤM LẠI BẢN SỬA (INTERNAL RE-GRADING - BƯỚC QUAN TRỌNG NHẤT):**
   - Hãy quên rằng bạn vừa sửa bài này. Hãy đóng vai một Giám khảo độc lập thứ 2 chấm lại bản 'annotated_essay' vừa tạo.
   - **Luật Nội dung (Content Rule):** Bản sửa chỉ sửa ngữ pháp/từ vựng, KHÔNG THỂ sửa lỗi thiếu số liệu/thiếu so sánh của bài gốc. Nếu bài gốc TA 6.0, bản sửa TA vẫn là 6.0 (hoặc tối đa 7.0 nếu diễn đạt rõ hơn).
   - **Kết luận:** Điểm 'revised_score' PHẢI là điểm thực tế của bản sửa, KHÔNG ĐƯỢC mặc định là 9.0.
Cấu trúc JSON:
```json
{
  "original_score": {
      "task_achievement": "Điểm TA của bài làm gốc (User's essay)",
      "cohesion_coherence": "Điểm CC của bài làm gốc",
      "lexical_resource": "Điểm LR của bài làm gốc",
      "grammatical_range": "Điểm GRA của bài làm gốc",
      "overall": "Điểm Overall của bài làm gốc (Average)"
  },
  "errors": [
    {
      "category": "Grammar" hoặc "Vocabulary",
      "type": "Tên Lỗi",
      "impact_level": "High" | "Medium" | "Low",
      "explanation": "Giải thích ngắn gọn lỗi.",
      "original": "đoạn văn bản sai",
      "correction": "đoạn văn bản đúng (VIẾT IN HOA)"
    }
  ],
  "annotated_essay": "Phiên bản bài làm đã được sửa lỗi (giữ nguyên cấu trúc các đoạn văn). Bọc từ sai trong thẻ <del>...</del> và từ sửa đúng trong thẻ <ins class='grammar'>...</ins> hoặc <ins class='vocab'>...</ins>. Nội dung sửa đúng phải viết IN HOA.",
   "revised_score": {
      "word_count_check": "BẮT BUỘC GHI SỐ TỪ CỦA BẢN SỬA (Ví dụ: '220 words - Too long')",
      "logic_re_evaluation": "Giải thích tại sao bị trừ điểm (Ví dụ: 'Dù sạch lỗi ngữ pháp nhưng bài viết dài 220 từ, vi phạm nguyên tắc súc tích, nên TA chỉ đạt 8.0').",
      "task_achievement": "Điểm TA thực tế (phạt nặng nếu dài dòng)",
      "cohesion_coherence": "Điểm CC",
      "lexical_resource": "Điểm LR",
      "grammatical_range": "Điểm GRA",
      "overall": "Điểm trung bình (Làm tròn theo Quy tắc làm tròn điểm bài viết theo chuẩn IELTS)"
          *   Làm tròn đến nửa band gần nhất (.0 hoặc .5).
          *   **NGOẠI LỆ BẮT BUỘC:**
              *   Điểm trung bình có đuôi **.25** -> BẮT BUỘC làm tròn **XUỐNG** số nguyên (Ví dụ: 8.25 -> 8.0).
              *   Điểm trung bình có đuôi **.75** -> BẮT BUỘC làm tròn **XUỐNG** .5 (Ví dụ: 8.75 -> 8.5).
  }
}
```
"""

# ==========================================
# 3. HELPER FUNCTIONS
# ==========================================

def clean_json(text):
    # Tìm đoạn văn bản nằm giữa dấu ngoặc nhọn { ... } đầu tiên và cuối cùng
    match = re.search(r"(\{[\s\S]*\})", text)
    if match:
        content = match.group(1).strip()
        # Loại bỏ các ký tự điều khiển lỗi
        content = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', content)
        return content
    return None

def parse_guide_response(text):
    j_str = clean_json(text)
    if not j_str: return None
    try:
        return json.loads(j_str)
    except:
        # Nếu lỗi JSON, thử quét tay các trường quan trọng (fallback)
        return {
            "task_type": "IELTS Task 1",
            "intro_guide": "Hãy paraphrase đề bài bằng từ đồng nghĩa.",
            "overview_guide": "Nêu xu hướng chung và đặc điểm nổi bật.",
            "body1_guide": "Mô tả chi tiết nhóm số liệu 1.",
            "body2_guide": "Mô tả chi tiết nhóm số liệu 2."
        }

def process_grading_response(full_text):
    """
    Hàm xử lý kết quả chấm điểm (CHUẨN TỪ APP CHẤM ĐIỂM).
    Tách biệt:
    1. Markdown Text (Phân tích chi tiết ở đầu).
    2. JSON Data (Điểm số và lỗi ở cuối).
    """
    json_str = clean_json(full_text)
    
    # Mặc định
    markdown_part = full_text
    data = {
        "errors": [], 
        "annotatedEssay": None, 
        "revisedScore": None, 
        "originalScore": {
            "task_achievement": "-", "cohesion_coherence": "-", 
            "lexical_resource": "-", "grammatical_range": "-", "overall": "-"
        }
    }
    
    if json_str:
        # Tách phần Markdown (trước JSON)
        markdown_part = full_text.split("```json")[0].strip()
        # Nếu AI không dùng code block, thử split bằng ký tự '{' đầu tiên của JSON
        if "original_score" in markdown_part: # Dấu hiệu JSON bị lẫn
             parts = full_text.split("{", 1)
             markdown_part = parts[0].strip()

        try:
            parsed = json.loads(json_str)
            data["errors"] = parsed.get("errors", [])
            data["annotatedEssay"] = parsed.get("annotated_essay")
            data["revisedScore"] = parsed.get("revised_score")
            data["originalScore"] = parsed.get("original_score", {})
        except json.JSONDecodeError:
            pass

    return markdown_part, data

# --- FILE EXPORT ---
def register_vietnamese_font():
    try:
        font_reg = "Roboto-Regular.ttf"
        font_bold = "Roboto-Bold.ttf"
        if not os.path.exists(font_reg):
            r = requests.get("https://github.com/googlefonts/roboto/raw/main/src/hinted/Roboto-Regular.ttf")
            with open(font_reg, "wb") as f: f.write(r.content)
        if not os.path.exists(font_bold):
            r = requests.get("https://github.com/googlefonts/roboto/raw/main/src/hinted/Roboto-Bold.ttf")
            with open(font_bold, "wb") as f: f.write(r.content)
        pdfmetrics.registerFont(TTFont('Roboto', font_reg))
        pdfmetrics.registerFont(TTFont('Roboto-Bold', font_bold))
        addMapping('Roboto', 0, 0, 'Roboto')
        addMapping('Roboto', 1, 0, 'Roboto-Bold')
        return True
    except: return False

def create_docx(data, topic, essay, analysis):
    doc = Document()
    doc.add_heading('IELTS ASSESSMENT REPORT', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('1. DETAILED ANALYSIS', level=1)
    doc.add_paragraph(analysis) # Phân tích chi tiết từ Markdown
    
    # Thêm bảng điểm
    doc.add_heading('2. SCORE BREAKDOWN', level=1)
    scores = data.get("originalScore", {})
    p = doc.add_paragraph()
    p.add_run(f"Overall Band: {scores.get('overall', '-')}\n").bold = True
    p.add_run(f"TA: {scores.get('task_achievement', '-')}, CC: {scores.get('cohesion_coherence', '-')}, LR: {scores.get('lexical_resource', '-')}, GRA: {scores.get('grammatical_range', '-')}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf(data, topic, essay, analysis):
    register_vietnamese_font()
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = [Paragraph("IELTS ASSESSMENT REPORT", styles['Title'])]
    
    # Analysis
    elements.append(Paragraph("DETAILED ANALYSIS", styles['Heading1']))
    # Clean markdown basic symbols for PDF
    safe_text = html.escape(analysis).replace('\n', '<br/>').replace('**', '').replace('#', '')
    elements.append(Paragraph(safe_text, styles['Normal']))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer

# ==========================================
# 4. UI: QUẢN LÝ TRẠNG THÁI (SESSION STATE)
# ==========================================
if "step" not in st.session_state: st.session_state.step = 1 
if "guide_data" not in st.session_state: st.session_state.guide_data = None
if "grading_result" not in st.session_state: st.session_state.grading_result = None
if "saved_topic" not in st.session_state: st.session_state.saved_topic = ""
if "saved_img" not in st.session_state: st.session_state.saved_img = None

# ==========================================
# 5. GIAO DIỆN CHÍNH (THEO YÊU CẦU MỚI)
# ==========================================

# TIÊU ĐỀ CHÍNH
st.markdown('<div class="main-header">🎓 IELTS Writing Task 1 – Examiner-Guided</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Learning & Scoring Based on IELTS Band Descriptors</div>', unsafe_allow_html=True)

if st.session_state.step == 1:
    
    # STEP 1 – Task 1 Question (Đã đổi lên trên)
    st.markdown('<div class="step-header">STEP 1 – Task 1 Question</div>', unsafe_allow_html=True)
    st.markdown('<div class="step-desc">Paste the official task question here</div>', unsafe_allow_html=True)
    question_input = st.text_area("Question", height=150, placeholder="The chart below shows...", key="q_input", label_visibility="collapsed")

    # STEP 2 – Visual Data (Đã đổi xuống dưới)
    st.markdown('<div class="step-header">STEP 2 – Visual Data </div>', unsafe_allow_html=True)
    st.markdown('<div class="step-desc">Upload chart / graph / table / diagram / map </div>', unsafe_allow_html=True)
    uploaded_image = st.file_uploader("Upload Image", type=['png', 'jpg', 'jpeg'], key="img_input", label_visibility="collapsed")
    
    img_data = None
    if uploaded_image:
        img_data = Image.open(uploaded_image)
        st.image(img_data, caption='Uploaded Visual Data', width=400)

    # STEP 3    
    st.markdown('<div class="step-header">STEP 3 – Examiner Workflow</div>', unsafe_allow_html=True)
    
    # --- PHẦN HTML NÀY PHẢI VIẾT SÁT LỀ TRÁI (KHÔNG THỤT DÒNG) ---
    workflow_html = """
<style>
    .wf-container { display: grid; grid-template-columns: 1fr 1fr; gap: 15px; margin-bottom: 20px; }
    .wf-card { background-color: #FFFFFF; border: 1px solid #E2E8F0; border-radius: 8px; padding: 15px; display: flex; align-items: center; }
    .wf-icon { width: 40px; height: 40px; background-color: #F0F9FF; color: #0284C7; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-size: 20px; margin-right: 15px; flex-shrink: 0; }
    .wf-title { font-weight: 700; font-size: 0.95rem; color: #1E293B; }
    .wf-desc { font-size: 0.85rem; color: #64748B; line-height: 1.4; }
</style>
<div class="wf-container">
    <div class="wf-card">
        <div class="wf-icon">🔍</div>
        <div class="wf-content">
            <div class="wf-title">1. Task Analysis</div>
            <div class="wf-desc">Analyze visual data to identify chart type.</div>
        </div>
    </div>
    <div class="wf-card">
        <div class="wf-icon">🧠</div>
        <div class="wf-content">
            <div class="wf-title">2. Data Grouping & Planning</div>
            <div class="wf-desc">Organise key features and trends logically..</div>
        </div>
    </div>
    <div class="wf-card">
        <div class="wf-icon">✍️</div>
        <div class="wf-content">
            <div class="wf-title">3. Guided Writing</div>
            <div class="wf-desc">Support writing with clear structure and useful vocabulary.</div>
        </div>
    </div>
    <div class="wf-card">
        <div class="wf-icon">⚖️</div>
        <div class="wf-content">
            <div class="wf-title">4. Band Score Evaluation</div>
            <div class="wf-desc">Evaluate based on official Band Descriptors.</div>
        </div>
    </div>
</div>
"""
    # GỌI LỆNH RENDER
    st.markdown(workflow_html, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)

    # Nút bấm xử lý (vẫn sử dụng question_input và img_data đã khai báo ở trên)
    if st.button("🔍 Analyze & Guide (Start Learning)", type="primary", width="stretch"):
        if not question_input or not img_data:
            st.warning("⚠️ Vui lòng nhập đầy đủ Đề bài và tải Ảnh lên để bắt đầu.")
        else:
            st.session_state.saved_topic = question_input
            st.session_state.saved_img = img_data
                   
            with st.spinner("🧠 The examiner is analysing the visual data and providing step-by-step guidance on how to write the answer..."):
                    # Prompt Tutor Vạn Năng: Tự động thích ứng theo từng dạng bài
                    prompt_guide = """
                    Bạn là một Siêu Giáo viên IELTS Writing (Band 9.0). Nhiệm vụ của bạn là phân tích hình ảnh đầu vào và viết hướng dẫn thực hành chi tiết.
                    **YÊU CẦU ĐẶC BIỆT (CHẾ ĐỘ PHÂN TÍCH KỸ):** Bạn không cần phải trả lời nhanh. Hãy dành thời gian "suy nghĩ" để phân tích thật sâu và chi tiết (Step-by-step Analysis).
                    # STRICT OUTPUT RULES (BẮT BUỘC TUÂN THỦ):
                    1.  **NO MARKDOWN LISTS:** Tuyệt đối KHÔNG được tự ý chuyển đổi định dạng sang gạch đầu dòng (bullet points) của Markdown.
                    2.  **HTML ONLY:** Output bắt buộc phải giữ nguyên các thẻ HTML: `<ul>`, `<li>`, `<b>`, `<br>`, `<code>`, `<div>`. Hệ thống chỉ render được HTML, nếu bạn dùng Markdown sẽ bị lỗi hiển thị.
                        
                    **BƯỚC 1: NHẬN DIỆN LOẠI BÀI (QUAN TRỌNG)**
                    Hãy nhìn hình ảnh và xác định nó thuộc loại nào:
                    1. **Change Over Time** (Line, Bar, Table, Pie có năm tháng): Cần từ vựng xu hướng (increase, decrease).
                    2. **Static Chart** (Pie, Table, Table 1 năm): Cần từ vựng so sánh (higher, lower, accounts for).
                    3. **Map (Bản đồ):** Cần từ vựng phương hướng (North, South) và sự thay đổi (demolished, constructed). Tuyệt đối không dùng "increase/decrease" cho nhà cửa.
                    4. **Process (Quy trình):** Cần câu Bị động (Passive voice) và từ nối trình tự (First, Then, Finally).
                    5. **Mixed (Kết hợp):** Cần hướng dẫn cách liên kết 2 biểu đồ.
                    
                    
                    **BƯỚC 2: SOẠN HƯỚNG DẪN (OUTPUT JSON)**

                    # =================================================================
                    # 🔴 TRƯỜNG HỢP 1: DẠNG "STATIC COMPARISON" (Pie, Bar, Table - 1 Năm/Không năm)
                    # (Tư duy cốt lõi: Ranking (Xếp hạng) & Proportion (Tỷ trọng))
                    # =================================================================

                    <br><i>(LƯU Ý: Nhìn Năm trong đề bài để quyết định THÌ cho toàn bài viết)</i>
                    <br>Tuyệt đối <b>KHÔNG</b> dùng dấu gạch ngang để chỉ khoảng số (VD: <i>7-14%</i>).
                    <br>👉 <b>Phải viết chữ:</b> <i>"between 7% and 14%"</i> hoặc <i>"from 7% to 14%"</i>.
                    1. **"intro_guide" (Paraphrase):**                    
    <ul>        
        <!-- CẤU TRÚC CHUẨN (CỐ ĐỊNH - KHÔNG ĐƯỢC THAY ĐỔI) -->
        <li><b>Cấu trúc chuẩn (Formula):</b> <code>[Subject] + [Finite Verb] + [Object/Topic] + [Place] + [Time (Specific Year)]</code>.</li>
    
        <li><b>Subject (Lưu ý quan trọng):</b>
            <br>- <b>Từ vựng bắt buộc:</b> Phải dùng các từ chỉ tỷ lệ như <i>The proportion of, The percentage of, The share of</i>.
            <br>- <b>Tuyệt đối tránh:</b> Không viết "The chart shows the immigration..." mà phải là "The chart shows the percentage of immigrants...".
            <br>- <b>Hòa hợp chủ ngữ:</b> <i>The pie charts compare...</i> (Số nhiều) hoặc <i>The pie chart gives information about...</i> (Số ít).</li>
    
        <li><b>Topic (The "What"):</b>
            <br>- <b>Cách 1 (Basic):</b> "...how people spent their money on different commodities..."
            <br>- <b>Cách 2 (Advanced):</b> "...the breakdown of expenses for..."
            <br>- <b>Cách 3 (Comparative):</b> "...the differences in the consumption of [Category]..."</li>
            
            <div style="background-color:#e8f5e9; padding:10px; border-radius:5px; border: 1px dashed #27ae60;">
                <strong style="color:#2ecc71;">🎓 2. EXAMINER'S EXTRA (Chủ đề Tiền tệ/Chi tiêu - Rất hay thi):</strong>
                <br><i>(Áp dụng khi đề bài là "Spending / Budget / Expenses")</i>
                <br>- <b>Cách 1:</b> "...how people <b>allocated their budget</b> to different commodities..."
                <br>- <b>Cách 2:</b> "...the <b>distribution of expenses</b> for..."
                <br>- <b>Cách 3:</b> "...consumption patterns regarding [Category]..."
            </div>
    
        <li><b>Verb (Động từ khuyên dùng):</b>
            <br><i>compares / gives information about / shows how many / shows the distribution of</i>.</li>
            
        <!-- PHẦN PARAPHRASE GỢI Ý THÊM -->
        <li><div style="background-color:#fff3e0; padding:10px; border-radius:5px; border: 1px dashed #f39c12; margin-top:5px;">
            <strong style="color:#d35400;">🔄 GỢI Ý PARAPHRASING (VOCABULARY INPUT):</strong>
            <br><i>(AI tự động liệt kê 3 từ hoặc cụm từ thay thế cho từ gốc trong đề bài)</i>
            <ul style="list-style-type: none; padding-left: 10px; margin-top: 5px; font-size: 0.9rem;">
               <li>🔹 <b>Topic (Chủ đề):</b> [Trích từ gốc] &rarr; <i>[AI điền từ thay thế]</i></li>
               <li>🔹 <b>Category (Đối tượng):</b> [Trích từ gốc] &rarr; <i>[AI điền từ thay thế]</i></li>
               <li>🔹 <b>Context (Phạm vi):</b> [Trích từ gốc] &rarr; <i>[AI điền từ thay thế]</i></li>
            </ul>
        </div></li>      
        <li><div style="background-color:#e6fffa; padding:10px; border-radius:5px; margin-top:5px; border-left: 4px solid #00b894;">
            <b>📝 Nội dung mẫu (Sample Intro):</b><br>                             
            <i>"The pie charts compare the percentage of people immigrating to a European country with the proportion of emigrants leaving that nation for various reasons in 2009."</i>
        </div></li>
    </ul>

                    2. **"overview_guide" (Ranking + Gap):**
                        - <ul>
                         <!-- ================================================================================== -->
                         <!-- PHẦN 1: KHO MẪU CÂU (CHẾ ĐỘ IN ẤN - KHÔNG TƯ DUY) -->
                         <!-- SYSTEM_ALERT: 
                              Dưới đây là DỮ LIỆU CỐ ĐỊNH (REFERENCE LIBRARY).
                              Nhiệm vụ của bạn: SAO CHÉP Y NGUYÊN TỪNG DÒNG HTML BÊN DƯỚI.
                              TUYỆT ĐỐI KHÔNG KHÔNG PHÂN TÍCH, KHÔNG ĐƯỢC LỌC. -->
                         <!-- ================================================================================== -->
                         <li>
                             <div style="background-color:#f8f9fa; border:1px solid #bdc3c7; border-radius:5px; padding:15px; margin-bottom:15px;">
                             <strong style="color:#c0392b;">⛔ LƯU Ý QUAN TRỌNG (EXAMINER'S NOTE):</strong>
                             <br><i>1. Tuyệt đối <b>KHÔNG</b> nhắc đến tên nhóm <b>"Other/Others"</b>.</i>
                             <br><i>2. <b>TRÁNH</b> khẳng định "X là thấp nhất" (The lowest) nếu trong biểu đồ còn có nhóm "Other" nhỏ hơn nó. Thay vào đó, hãy dùng từ mang nghĩa "nhóm nhỏ/thiểu số".</i>
                                                          
                             <hr style="border-top: 1px dashed #ccc; margin: 15px 0;">
                             <strong style="color:#d35400;">📚 KHO MẪU CÂU OVERVIEW (STATIC):</strong>
                            <br><i>(Tập trung vào Cái lớn nhất & Sự chênh lệch)</i>
                             
                             <!-- KHỐI 1: TỔNG QUÁT -->
                             <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
                                 <b>► 1. Cấu trúc tổng quát:</b>
                                 <!-- CẤU TRÚC CHUẨN (CỐ ĐỊNH - KHÔNG ĐƯỢC THAY ĐỔI) -->
                                 <br><code>Overall, it is clear/noticeable that &#91;Most Popular Category&#93; accounts for the largest share. In contrast, &#91;Least Popular Category&#93; makes up the smallest proportion.</code>
                             </div>

                             <!-- KHỐI 2: FEATURE 1 (DOMINANT CATEGORY) -->
                             <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
                                 <b>► 2. Các mẫu câu tả Nhóm Lớn Nhất:</b>
                                 <br>✅ <b>Loại 1: Áp đảo tuyệt đối</b>
                                 <br>"<b>&#91;Category A&#93;</b> is by far the most popular option/reason."
                                 <br>✅ <b>Loại 2: Chiếm đa số</b>
                                 <br>"The majority of <b>&#91;Topic&#93;</b> is allocated to <b>&#91;Category A&#93;</b>."                           
                                 <br>✅ <b>Loại 3: So sánh đối lập (2 biểu đồ)</b>
                                 <br>"While <b>&#91;Category A&#93;</b> is the dominant figure in <b>&#91;Group 1&#93;</b>, <b>&#91;Category B&#93;</b> takes the lead in <b>&#91;Group 2&#93;</b>."
                             </div>

                             <!-- KHỐI 3: FEATURE 2 (OTHER FEATURES) -->
                             <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
                                 <b>► 3. Các mẫu câu tả Đặc điểm phụ:</b>                             
                                 <br>✅ <b>Loại 1: Nhóm thấp nhất (Có tên cụ thể)</b>
                                 <br>"<b>&#91;Category C&#93;</b> represents the least significant portion among the specified categories."                              
                                 <br>✅ <b>Loại 2: Điểm tương đồng (Similarity)</b>
                                 <br>"It is also noticeable that <b>&#91;Category B&#93;</b> accounts for a significant portion in both charts."
                             </div>
                         </li>
                         
                         <hr style="border-top: 1px dashed #ccc; margin: 15px 0;">

                         <!-- PHẦN 2: PHÂN TÍCH -->
                         <li>
                             <b>🔍 PHÂN TÍCH BÀI NÀY (Selection & Drafting):</b>
                             <br><b>1. Xác định "Hạng mục Phổ biến nhất" (The Dominant Category):</b>
                             <br>- Nhìn vào biểu đồ/bảng: <b>Hạng mục (Category)</b> nào có số liệu cao nhất/lớn nhất?
                             <br><i>(Lưu ý: Phải là hạng mục có tên cụ thể, không chọn 'Total' hay 'Other')</i>.
                             <br>- <b>Quyết định:</b> Chọn Mẫu nào?
                             <br>👉 <b>Viết nháp Câu 1:</b> <i>[AI viết câu hoàn chỉnh dựa trên Mẫu đã chọn + Chia đúng thì]</i>
                             <br>
                              <br><b>2. Xác định "Hạng mục Đối lập/Đặc biệt" (Contrast/Exception):</b>
                             <br>- <b>Tìm sự chênh lệch (Gap):</b> Khoảng cách giữa Hạng mục cao nhất và thấp nhất có lớn không?
                             <br>- <b>Tìm Hạng mục thấp nhất:</b> Cái nào thấp nhất (trừ Other)?
                             <br>- <b>Tìm điểm chung (Similarity):</b> Có hạng mục nào cao/thấp đồng đều ở tất cả các nhóm không?
                             <br>- <b>Quyết định:</b> Chọn Mẫu nào?
                             <br>👉 <b>Viết nháp Câu 2:</b> <i>[AI viết câu hoàn chỉnh dựa trên Mẫu đã chọn + Chia đúng thì]</i>
                         </li>

                         <!-- PHẦN 3: TỪ VỰNG -->
                         <li><b>🔑 TỪ VỰNG GỢI Ý (Vocabulary):</b>
                             <br>- <b>Verbs (Chiếm bao nhiêu):</b> <i>account for / make up / constitute / represent / comprise</i>.
                             <br>- <b>Adjectives (So sánh):</b> <i>dominant / significant / negligible (không đáng kể) / overwhelming</i>.
                             <br>- <b>Nouns:</b> <i>The vast majority / a small fraction / the lion's share</i>.
                         </li>

                         <!-- PHẦN 4: BÀI MẪU -->
                         <li><div style="background-color:#fff3e0; padding:15px; border-radius:8px; margin-top:10px; border-left: 5px solid #ff9f43;">
                             <b>📝 Nội dung mẫu (Sample Overview):</b><br>
                             <div style="margin-top:5px; font-style: italic; color: #5d4037;">
                             [AI hãy ghép 2 câu nháp ở "PHẦN 2: PHÂN TÍCH" và thành đoạn Overview hoàn chỉnh.]
                             </div>
                         </div></li>
                       </ul>

                    3. **"body1_guide" (Thân bài 1 - Nhóm Lớn Nhất / Nổi Bật Nhất):**
    <ul>
     <!-- ================================================================================== -->
     <!-- MA TRẬN ĐA DẠNG HÓA CẤU TRÚC (GRA BOOSTER) -->
     <!-- ================================================================================== -->
     <li><div style="background-color:#e8f5e9; padding:15px; border:1px solid #4caf50; border-radius:5px; margin-bottom:15px;">
         <strong style="color:#2e7d32;">🎨 KỸ THUẬT ĐA DẠNG HÓA CHỦ NGỮ (SUBJECT VARIATION):</strong>
         <br><i>(Examiner yêu cầu: Không được viết 2 câu liên tiếp có cùng kiểu chủ ngữ. Hãy xoay tua 4 loại sau:)</i>
         
         <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin-top: 10px; font-size: 0.9rem;">
             <!-- LOẠI 1 -->
             <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #a5d6a7;">
                 <b>1. Category làm chủ ngữ:</b>
                 <br><i>"<b>Gas</b> accounted for 30%..."</i>
             </div>
             <!-- LOẠI 2 -->
             <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #a5d6a7;">
                 <b>2. Từ chỉ Lượng làm chủ ngữ:</b>
                 <br><i>"<b>The proportion of Gas</b> was 30%..."</i>
             </div>
             <!-- LOẠI 3 -->
             <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #a5d6a7;">
                 <b>3. Xếp hạng làm chủ ngữ:</b>
                 <br><i>"<b>The leading source</b> was Gas..."</i>
             </div>
             <!-- LOẠI 4 -->
             <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #a5d6a7;">
                 <b>4. Nơi chốn/Người làm chủ ngữ:</b>
                 <br><i>"<b>France</b> produced 30%..."</i>
             </div>
         </div>
     </div></li>
     <!-- ================================================================================== -->
     <!-- PHẦN 1: KHO MẪU CÂU STATIC (REFERENCE LIBRARY) -->
     <!-- SYSTEM_ALERT: 
          Bắt buộc hiển thị đủ các Cấu trúc so sánh, Cấu trúc Math (Gấp đôi/Gấp ba) và Cấu trúc mệnh đề quan hệ (Which/With).
          Dưới đây là DỮ LIỆU CỐ ĐỊNH (REFERENCE LIBRARY).
          Nhiệm vụ của bạn: SAO CHÉP Y NGUYÊN TỪNG DÒNG HTML BÊN DƯỚI.
          TUYỆT ĐỐI KHÔNG ĐƯỢC LỌC. -->
     <!-- ================================================================================== -->
     <li>     
         <div style="background-color:#f8f9fa; border:1px solid #bdc3c7; border-radius:5px; padding:15px; margin-bottom:15px;">
         <strong style="color:#d35400;">📚 KHO MẪU CÂU BODY 1 (STATIC & COMPARISON):</strong>
         <br><i>(Áp dụng kỹ thuật: Ranking + Math Language + Complex Clauses)</i>
         <br><i>(Lưu ý: <b>[V_Tense]</b> = Chia động từ theo Năm của đề bài (Quá khứ/Hiện tại/Tương lai))</i>
         
         <!-- KHỐI 1: CÂU MỞ ĐẦU (STARTING WITH THE HIGHEST) -->
         <!-- KEEP_FULL_BLOCK_1 -->
         <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
             <b>► Bước 1: Câu mở đầu (Top Ranking):</b>
             <br><i>(Luôn bắt đầu bằng Category lớn nhất)</i>
             
             <br>✅ <b>Mẫu 1: Dùng động từ "Account for/Constitute":</b>
             <br>"<b>Regarding &#91;Category A&#93;</b>, it <b>accounts for</b> the largest proportion of <b>&#91;Total&#93;</b>, at <b>&#91;Data&#93;</b>."
             
             <br>✅ <b>Mẫu 2: Nhấn mạnh vị trí số 1 (Page 5/6):</b>
             <br>"<b>&#91;Category A&#93;</b> <b>ranks first</b> among all reasons, with <b>&#91;Data&#93;</b>."
             <br>"<b>&#91;Category A&#93;</b> <b>takes the lead</b>, comprising <b>&#91;Data&#93;</b> of the total."
             
             <br>✅ <b>Mẫu 3: Mệnh đề quan hệ rút gọn (Advanced):</b>
             <br>"Standing at <b>&#91;Data&#93;</b>, <b>&#91;Category A&#93;</b> is the most popular choice."
         </div>

         <!-- KHỐI 2: SO SÁNH -->
         <!-- KEEP_FULL_BLOCK_2 -->
         <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
             <b>► Bước 2: So sánh gấp lần (Math Language):</b>
             <br><i>(Dùng để so sánh Category A với Category B)</i>
             <br>✅ <b>Cấu trúc Gấp đôi/Ba (Double/Triple - Page 4):</b>
             <br>"The figure for <b>&#91;Category A&#93;</b> is <b>exactly/nearly double</b> that of <b>&#91;Category B&#93;</b> (Data A vs Data B)."
             <br>"<b>&#91;Category A&#93;</b> is <b>three times as high as</b> <b>&#91;Category B&#93;</b>."
             
             <br>✅ <b>Cấu trúc "Which" clause:</b>
             <br>"<b>&#91;Category A&#93;</b> stands at <b>&#91;Data&#93;</b>, <b>which is significantly higher than</b> the figure for <b>&#91;Category B&#93;</b>."
         </div>

         <!-- KHỐI 3: ĐƯA SỐ LIỆU PHỨC HỢP -->
         <!-- KEEP_FULL_BLOCK_3 -->
         <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
             <b>► Bước 3: Đưa số liệu phức hợp:</b>
             <br>✅ <b>Cấu trúc "With respective figures":</b>
             <br>"<b>&#91;Cat A&#93;</b> and <b>&#91;Cat B&#93;</b> are the top choices, <b>with respective figures being</b> <b>&#91;Data A&#93;</b> and <b>&#91;Data B&#93;</b>."
             <br>✅ <b>Cấu trúc "Making it":</b>
             <br>"<b>&#91;Data&#93;</b> of people chose <b>&#91;Cat A&#93;</b>, <b>making it the most common</b> reason cited."
         </div>
         </div>
     </li>
     
     <hr style="border-top: 1px dashed #ccc; margin: 15px 0;">

     <!-- ================================================================================== -->
     <!-- PHẦN 2: THỰC HÀNH TƯ DUY & LẮP RÁP -->
     <!-- ================================================================================== -->
     <li>     
         <b>✍️ THỰC HÀNH TƯ DUY & LẮP RÁP:</b>
         <br><b>🔍 Phân tích Dữ liệu Body 1:</b>
         <br>- Chọn nhóm: Các hạng mục LỚN NHẤT (Major Categories).
         <br>- Mục tiêu: Mô tả số liệu và so sánh độ lớn.
         <br>         
         <br><b>✅ BƯỚC 1: Xử lý Top 1 (The Winner)</b>
         <br>- <b>Quyết định:</b> Chọn Mẫu nào? Dùng Chủ ngữ loại mấy (1, 2, 3 hay 4)?
         <br>👉 <b>Viết nháp Câu 1:</b> <i>[AI viết câu hoàn chỉnh dựa trên Mẫu đã chọn + Chia đúng thì]</i>
         <br>
         <br><b>✅ BƯỚC 2: So sánh với Top 2 (The Runner-up)</b>
         <br>- <b>Logic Toán học:</b> Top 1 có gấp đôi Top 2 không? Hay chỉ hơn một chút?
         <br>- <b>Quyết định:</b> Chọn Mẫu nào? Dùng Chủ ngữ loại mấy (1, 2, 3 hay 4)?
         <br>👉 <b>Viết nháp Câu 2:</b> <i>[AI viết câu hoàn chỉnh dựa trên Mẫu đã chọn + Chia đúng thì]</i>
         <br>
         <br><b>✅ BƯỚC 3: Gom nhóm (nếu có Top 3 lớn)</b>
         <br>- <b>Quyết định:</b> Chọn Mẫu nào? Dùng Chủ ngữ loại mấy (1, 2, 3 hay 4)?
         <br>👉 <b>Viết nháp Câu 3:</b> <i>[AI viết câu chốt số liệu]</i>
     </li>

     <!-- ================================================================================== -->
     <!-- PHẦN 3: BỘ CÔNG CỤ NÂNG BAND -->
     <!-- ================================================================================== -->
     <li>     
     <strong style="color:#d35400;">🔑 TỪ VỰNG & NGỮ PHÁP "STATIC" (VOCABULARY BANK):</strong>
         <br>
         <ul style="margin-top:5px;">
            <!-- NHÓM 1: ĐỘNG TỪ CHỈ TỶ TRỌNG -->
            <li><b>1. Động từ chiếm tỷ trọng (Percentage Verbs):</b>
                <br><i>account for / constitute / comprise / represent / make up</i> (+ %).
            </li>

            <!-- NHÓM 2: TỪ VỰNG XẾP HẠNG (RANKING - PAGE 5) -->
            <li><b>2. Ranking Vocabulary:</b>
                <br>- <i>Rank first / second / third.</i>
                <br>- <i>Take the lead / Is the leading factor.</i>
                <br>- <i>The most popular / common / dominant.</i>
            </li>

            <!-- NHÓM 3: TỪ VỰNG SO SÁNH SỐ LIỆU (MATH - PAGE 4) -->
            <li><b>3. Mathematical Comparisons:</b>
                <br>- <i>Two times as high as / Three times as much as.</i>
                <br>- <i>Double / Triple (Dùng như động từ hoặc tính từ).</i>
                <br>- <i>Half of / A quarter of.</i>
            </li>

             <!-- 4. CẤU TRÚC NGỮ PHÁP CAO CẤP (PAGE 5/6) -->
             <li><b>4. Advanced Grammar:</b>
                 <br>- <b>Making it:</b> <i>..., making it the largest category.</i>
                 <br>- <b>With respective figures:</b> <i>..., with respective figures being X and Y.</i>
                 <br>- <b>Compared to:</b> <i>X is high, compared to only 5% of Y.</i>
             </li>
         </ul>
     </li>

     <!-- PHẦN 4: BÀI MẪU -->
     <li><div style="background-color:#fff8e1; padding:15px; border-radius:8px; margin-top:10px; border-left: 5px solid #ffa502;">
         <b>📝 Nội dung mẫu (Sample Body 1):</b><br>
         <div style="margin-top:5px; font-style: italic; color: #5d4037;">
         [AI hãy ghép các câu nháp ở "PHẦN 2: THỰC HÀNH TƯ DUY & LẮP RÁP" thành đoạn Body 1 hoàn chỉnh.]
         </div>
     </div></li>
   </ul>

                    4. **"body2_guide" (Thân bài 2 - Nhóm Nhỏ / Nhóm Còn lại):**
    <ul>
     <!-- ================================================================================== -->
     <!-- PHẦN 1: KHO MẪU CÂU (REFERENCE LIBRARY) 
     <!-- SYSTEM_ALERT:           
          Dưới đây là DỮ LIỆU CỐ ĐỊNH (REFERENCE LIBRARY).
          Nhiệm vụ của bạn: SAO CHÉP Y NGUYÊN TỪNG DÒNG HTML BÊN DƯỚI.
          TUYỆT ĐỐI KHÔNG ĐƯỢC LỌC. -->
     <!-- ================================================================================== -->
     <li>
         <div style="background-color:#f8f9fa; border:1px solid #bdc3c7; border-radius:5px; padding:15px; margin-bottom:15px;">
         <strong style="color:#2980b9;">📚 KHO MẪU CÂU BODY 2 (REMAINING CATEGORIES):</strong>
         <br><i>(Tập trung vào sự tương phản và gom nhóm nhỏ)</i>
         
         <!-- KHỐI 1: CÂU CHUYỂN ĐOẠN -->
         <!-- KEEP_FULL_BLOCK_1 -->
         <div style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; margin-top:5px;">
             <b>► Bước 1: Câu chuyển đoạn (Transition):</b>
             <br><i>(Chọn 1 trong 2 mẫu sau)</i>
             <br>✅ <b>Mẫu 1 (Chuyển hướng):</b> "<b>Turning to</b> the remaining categories,..."
             <br>✅ <b>Mẫu 2 (Đối lập):</b> "<b>In contrast / By comparison</b>, the figures for <b>&#91;Category C & D&#93;</b> <b>[V_Tense: were/are]</b> significantly lower."
         </div>

         <!-- KHỐI 2: MIÊU TẢ NHÓM SỐ LIỆU NHỎ -->
         <!-- KEEP_FULL_BLOCK_2 -->
         <div style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; margin-top:5px;">
             <b>► Bước 2: Gom nhóm (Grouping):</b>
             <br><i>(Chọn 1 trong 2 mẫu sau)</i>
             <br>✅ <b>Mẫu 1 (Collectively):</b> "<b>&#91;Category C&#93;</b> and <b>&#91;Category D&#93;</b> <b>collectively [V_Tense: accounted/account] for</b> only <b>&#91;Total %&#93;</b>."
             <br>✅ <b>Mẫu 2 (Range - Chú ý Formal):</b> "The figures for X and Y <b>[V_Tense: were/are]</b> relatively small, <b>ranging from</b> <b>&#91;Data 1&#93;</b> to <b>&#91;Data 2&#93;</b>."
         </div>

          <!-- KHỐI 3: CẤU TRÚC "NEGLIGIBLE" -->
         <!-- KEEP_FULL_BLOCK_3 -->
         <div style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; margin-top:5px;">
             <b>► Bước 3: Mô tả số liệu rất nhỏ:</b>
             <br><i>(Chọn 1 trong 2 mẫu sau)</i>
             <br>✅ <b>Mẫu 1 (Negligible):</b> "The proportion of <b>&#91;Category E&#93;</b> <b>[V_Tense: was/is]</b> <b>negligible/insignificant</b>, at only <b>&#91;Data&#93;</b>."
             <br>✅ <b>Mẫu 2 (Minority):</b> "Only a <b>small minority</b> of people (<b>&#91;Data&#93;</b>) <b>[V_Tense: chose/choose]</b> <b>&#91;Category E&#93;</b>."
         </div>
         </div>
     </li>
     
     <hr style="border-top: 1px dashed #ccc; margin: 15px 0;">

     <!-- ================================================================================== -->
     <!-- PHẦN 2: THỰC HÀNH TƯ DUY & LẮP RÁP (PROCESS) -->
     <!-- ================================================================================== -->
     <li>     
         <b>✍️ THỰC HÀNH TƯ DUY & LẮP RÁP:</b>
         <br><b>🔍 Phân tích Dữ liệu Body 2:</b>
         <br>- Nhóm này gồm: Các hạng mục còn lại (nhỏ/thấp).
         <br>- <b>⚠️ REMINDER:</b> Kiểm tra lại thì (Quá khứ/Hiện tại) và cách viết khoảng số (from...to...).
         <br>
         <br><b>✅ BƯỚC 1: Xử lý Câu chuyển đoạn & Đối lập</b>
         <br>- <b>Quyết định:</b> Chọn Mẫu nào?
         <br>👉 <b>Viết nháp Câu 1:</b> <i>[AI viết câu chuyển + Chia đúng thì]</i>
         <br>
         <br><b>✅ BƯỚC 2: Xử lý các hạng mục tương đồng/nhỏ</b>
         <br>- <b>Quyết định:</b> Chọn Mẫu nào?
         <br>👉 <b>Viết nháp Câu 2:</b> <i>[AI viết câu mô tả nhóm nhỏ + Chia đúng thì]</i>
         <br>
         <br><b>✅ BƯỚC 3: Xử lý phần còn lại (Negligible)</b>
         <br>- <b>Quyết định:</b> Chọn Mẫu nào?
         <br>👉 <b>Viết nháp Câu 3:</b> <i>[AI viết câu chốt nhóm thấp nhất + Chia đúng thì]</i>
     </li>

     <!-- ================================================================================== -->
     <!-- PHẦN 3: GIẢI MÃ TỪ VỰNG -->
     <!-- ================================================================================== -->
     <li>
         <strong style="color:#d35400;">🔑 TỪ VỰNG & NGỮ PHÁP "ĂN ĐIỂM" (VOCABULARY BANK):</strong>
         <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 15px; margin-top: 10px; background-color:#fffcf5; padding:15px; border-radius:8px; border: 1px solid #f39c12;">
            <!-- CỘT 1 -->
            <div>
                <b style="color:#e67e22;">1. Từ nối & Đối lập:</b>
                <ul style="margin:5px 0; padding-left:20px; font-size:0.9rem;">
                    <li><b>Conversely / In contrast:</b> Ngược lại.</li>
                    <li><b>As for... / Regarding...:</b> Đối với...</li>
                    <li><b>A similar pattern:</b> Một kiểu mẫu tương tự.</li>
                </ul>
            </div>

            <!-- CỘT 2 -->
            <div>
                <b style="color:#e67e22;">2. Từ vựng nhóm nhỏ:</b>
                <ul style="margin:5px 0; padding-left:20px; font-size:0.9rem;">
                    <li><b>A negligible amount:</b> Một lượng không đáng kể.</li>
                    <li><b>Collectively:</b> Gộp chung lại.</li>
                    <li><b>The remainder:</b> Phần còn lại.</li>
                </ul>
            </div>
         </div>
     </li>

     <!-- ================================================================================== -->
     <!-- PHẦN 4: BÀI MẪU HOÀN CHỈNH -->
     <!-- ================================================================================== -->
     <li><div style="background-color:#e1f5fe; padding:15px; border-radius:8px; margin-top:10px; border-left: 5px solid #03a9f4;">
         <b>📝 NỘI DUNG MẪU (SAMPLE BODY 2):</b><br>
         <div style="margin-top:5px; font-style: italic; color: #5d4037;">
         [AI hãy ghép các câu nháp ở "PHẦN 2: THỰC HÀNH TƯ DUY & LẮP RÁP" thành đoạn Body 2 hoàn chỉnh.
         </div>
     </div></li>
   </ul>
                    # =================================================================
                    # 🔵 TRƯỜNG HỢP 2: DẠNG "CHANGE OVER TIME" (Line, Bar, Table, Pie nhiểu năm)
                    # (Tư duy cốt lõi: Trend (Xu hướng) & Speed (Tốc độ thay đổi))
                    # =================================================================

                    1. **"intro_guide" (Paraphrase):**                    
    <ul>
        <li><b>Cấu trúc chuẩn (Formula):</b> <code>[Subject] + [Finite Verb] + [Object/Topic] + [Place] + [Time]</code>.</li>
    
        <li><b>Subject (Lưu ý quan trọng):</b>
            <br>- <b>Xác định đúng chủ thể:</b> <i>[Xác định chính xác cái gì thay đổi]</i>.
            <br>- <b>Đơn vị trong bài này là:</b> <i>[Điền đơn vị cụ thể của bài, VD: million dollars / tonnes / %]</i>.
            <br>- <b>Tuyệt đối không đưa đơn vị tính vào chủ ngữ.</b> Ví dụ: Không viết <i>"The chart shows [Đơn vị của bài]..."</i> mà phải viết <i>"The chart shows the amount/number/proportion of..."</i>.
            <br>- <b>Hòa hợp chủ ngữ - động từ:</b> Nếu 1 biểu đồ dùng <i>shows/illustrates</i>. Nếu nhiều biểu đồ dùng <i>show/illustrate</i>.</li>
    
        <li><b>Cách đổi Chủ ngữ & Topic (The "What") cho bài này:</b>
            <br>- <b>Từ vựng gốc trong đề:</b> "<i>[Trích cụm từ gốc trong đề bài]</i>"
            <br>- <b>Gợi ý Paraphrase 1:</b> <i>[Viết phương án paraphrase 1. VD: The amount of money spent on...]</i>
            <br>- <b>Gợi ý Paraphrase 2:</b> <i>[Viết phương án paraphrase 2. VD: How much money was allocated to...]</i>
            <br><i>(Lưu ý: chọn từ Spending/Number/Percentage phù hợp).</i></li>
    
        <li><b>Verb (Động từ khuyên dùng):</b>
            <br><i>illustrates / gives information about / compares the data on / presents information about</i>.</li>
    
        <li><b>Time Paraphrase (Thời gian: [Năm đầu] - [Năm cuối]):</b>
            <br>- Cách 1: <i>Between [Năm đầu] and [Năm cuối]</i>.
            <br>- Cách 2: <i>Over a period of [Số năm] years starting from / commencing in [Năm đầu]</i>.</li>    
                
        <li><div style="background-color:#e6fffa; padding:10px; border-radius:5px; margin-top:5px; border-left: 4px solid #00b894;">
            <b>📝 Nội dung mẫu (Sample Intro):</b><br>                             
            <i>"[Viết câu Introduction theo hướng dẫn đã phân tích]"</i>
        </div></li>
    </ul>

                    2. **"overview_guide" (Trend + Ranking):**
                        - <ul>
                         <!-- ================================================================================== -->
                         <!-- PHẦN 1: KHO MẪU CÂU (CHẾ ĐỘ IN ẤN - KHÔNG TƯ DUY) -->                         
                         <!-- ================================================================================== -->
                         <li>
                             <div style="background-color:#f8f9fa; border:1px solid #bdc3c7; border-radius:5px; padding:15px; margin-bottom:15px;">
                             <strong style="color:#d35400;">📚 KHO MẪU CÂU:</strong>
                                                         
                             <!-- KHỐI 1: TỔNG QUÁT -->
                             <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
                                 <b>► 1. Cấu trúc tổng quát:</b>
                                 <br><code>Overall, &#91;Sentence 1: Trends&#93;. In addition, &#91;Sentence 2: Highlights&#93;.</code>
                             </div>
                                                        
                             <div style="background-color:#ffebee; border:1px solid #ef5350; padding:10px; margin-bottom:10px; border-radius:4px;">
                                 <strong style="color:#c62828;">⛔ LƯU Ý VỀ ĐỘNG TỪ (VERB RESTRICTIONS):</strong>
                                 <br>1. Với chủ ngữ là <b>SỐ LIỆU</b> (The figure / The percentage):
                                 <br>• <b>Nên dùng:</b> <i>increased, decreased, rose, fell</i> (Trực tiếp).
                                 <br>• <b>Chấp nhận:</b> <i>experienced, underwent</i> (Trải qua).
                                 <br>• <b>HẠN CHẾ:</b> <i>saw, witnessed</i> (Vì con số không có mắt để nhìn).
                                 
                                 <br>2. Với chủ ngữ là <b>HẠNG MỤC / NƠI CHỐN</b> (The UK / Farming):
                                 <br>👉 <b>Phải dùng:</b> <i>saw / witnessed / experienced ... <b>IN</b> ...</i>
                             </div>
                             
                             <!-- KHỐI 2: TRENDS (ĐÃ KIỂM TRA KỸ LƯỠNG - BUG FREE) -->
                             <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
                                 <b>► 2. Các mẫu câu Xu hướng (Trends):</b>
                                 <br><i>(Phân tích đề bài và hình ảnh để chọn Trends loại nào?, và chỉ hiển thị công thức)</i>                                
                                 
                                 <br>✅ <b>Loại 1: Đồng loạt Tăng/Giảm (Same Direction)</b>
                                 <br><i>(Dùng khi tất cả cùng Tăng hoặc cùng Giảm)</i>                                 
                                 <br>• <b>Số liệu:</b> "It is clear that the <b>total [number/amount/percentage] of &#91;Topic&#93;</b> increased/decreased <b>over the period shown</b>."
                                 <br>• <b>Hạng mục:</b> "It is noticeable that <b>&#91;Place/Category&#93;</b> witnessed an upward/downward trend <b>IN</b> <b>&#91;Topic&#93;</b> <b>throughout the given period</b>."
                                 
                                 <br>✅ <b>Loại 2: Xu hướng ngược (Mix / Opposite)</b>
                                 <br><i>(CHỈ DÙNG khi A Tăng còn B Giảm - Ngược chiều hoàn toàn)</i>                                 
                                 <br>• <b>Số liệu:</b> "It is clear that while the <b>figures for &#91;Category A&#93;</b> increased, the opposite was true for <b>&#91;Category C&#93;</b> <b>over the given period</b>."
                                 <br>• <b>Hạng mục:</b> "It is noticeable that while <b>&#91;Category A&#93;</b> saw an upward trend <b>IN</b> <b>&#91;Topic&#93;</b>, <b>&#91;Category C&#93;</b> experienced a decline <b>during the period shown</b>."                                 
                                 
                                 <br>✅ <b>Loại 3: Ngoại lệ (Exception)</b>
                                 <br><i>(Dùng khi đa số Tăng, chỉ có 1 cái Giảm/Ổn định)</i>                                 
                                 <br>• <b>Mẫu cơ bản:</b> "The figures for most categories increased, <b>with the exception of &#91;Category C&#93;</b>, <b>over the given period</b>."
                                 <br>• <b>Mẫu nâng cao:</b> "<b>With the exception of &#91;Category C&#93;</b>, all other categories <b>witnessed</b> an upward/downward trend <b>over the period shown</b>."
                                 
                                 <br>✅ <b>Loại 4: Ổn định/Dao động (Stability/Fluctuation)</b>
                                 <br><i>(Dùng khi có đường đi ngang hoặc dao động mạnh)</i>
                                 <br>"It is clear that while <b>&#91;Category A&#93;</b> changed significantly, the figure for <b>&#91;Category B&#93;</b> remained relatively stable <b>over the given period</b>."                                
                             </div>

                             <!-- KHỐI 3:  KEY FEATURES / HIGHLIGHTS (LIỆT KÊ ĐỦ 3 LOẠI) -->
                             <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
                                 <b>► 3. Các mẫu câu Điểm nổi bật (Key Features / Highlights):</b>
                                 <br><i>(Phân tích đề bài và hình ảnh để chọn Highlights loại nào?, và chỉ hiển thị công thức)</i>                               
                                 <br>✅ <b>Loại 1: Cao nhất/Thấp nhất (Ranking)</b>
                                 <br><i>(Dùng khi có 1 đường luôn nằm trên hoặc nằm dưới các đường khác)</i>
                                 <br>"<b>&#91;Category A&#93;</b> consistently had the highest figures throughout the period."                                 
                                 <br>✅ <b>Loại 2: Biến động lớn nhất (Biggest Change)</b>
                                 <br><i>(Dùng khi có 1 đường tăng/giảm mạnh nhất so với bọn còn lại)</i>
                                 <br>"<b>&#91;Category B&#93;</b> witnessed the most dramatic change."                               
                                 <br>✅ <b>Loại 3: Soán ngôi (Ranking Shift)</b>
                                 <br><i>(Dùng khi các đường cắt nhau)</i>
                                 <br>"The <b>figure for &#91;Category A&#93;</b> overtook <b>that of &#91;Category B&#93;</b> to become the dominant category."                               
                             </div>
                         </li>
                         
                         <hr style="border-top: 1px dashed #ccc; margin: 15px 0;">

                         # ==================================================================================
                         # PHẦN 2: CHIẾN THUẬT CHỌN MẪU (TEMPLATE MATCHING)
                         # Mục tiêu: Dạy học sinh nhìn hình -> Chọn mẫu -> Giải thích tại sao.
                         # ==================================================================================
                         <li>
                             <div style="background-color:#e1f5fe; padding:15px; border-radius:8px; border-left: 5px solid #0288d1; margin-bottom: 15px;">
                                 <strong style="color:#01579b; font-size: 1.1rem;">🔍 PHÂN TÍCH BÀI NÀY (Selection & Drafting):</strong>
                                 <br><i>(Học sinh nhìn hình để chọn Mẫu cho phù hợp)</i>                                 
                                 <br><br><b>1. Với Xu hướng (Trend - Câu 1):</b>
                                 <ul style="margin-top:5px; list-style-type: none; padding-left: 10px;">
                                     <li>👀 <b>Quan sát (Observation):</b> <i>[AI MÔ TẢ NGẮN GỌN DẤU HIỆU HÌNH ẢNH. Ví dụ: "Đường A đi lên mạnh, trong khi đường B lại đi xuống."]</i></li>
                                     <li>🧩 <b>Khớp với Mẫu số:</b> <b>[AI ĐIỀN LOẠI. Ví dụ: Loại 2 (Xu hướng ngược/Mix)]</b>.</li>
                                     <li>💡 <b>Lý do chọn:</b> <i>[AI GIẢI THÍCH LOGIC. Ví dụ: "Vì hai đường đi ngược chiều nhau nên bắt buộc dùng cấu trúc tương phản 'While/In contrast'."]</i></li>
                                 </ul>
                                 <br><b>2. Với Điểm nổi bật (Highlight - Câu 2):</b>
                                  <ul style="margin-top:5px; list-style-type: none; padding-left: 10px;">
                                     <li>👀 <b>Quan sát (Observation):</b> <i>[AI MÔ TẢ. Ví dụ: "Đường A luôn nằm cao nhất, không bị ai vượt qua."]</i></li>
                                     <li>🧩 <b>Khớp với Mẫu số:</b> <b>[AI ĐIỀN LOẠI. Ví dụ: Loại 1 (Cao nhất/Ranking)]</b>.</li>
                                     <li>💡 <b>Lý do chọn:</b> <i>[AI GIẢI THÍCH. Ví dụ: "Vì đường này luôn áp đảo (dominant) suốt giai đoạn nên dùng từ 'consistently the highest'."]</i></li>
                                 </ul>
                             </div>
                             
                             <!-- Phần Drafting: Tách riêng ra để học sinh thấy kết quả sau khi tư duy -->
                             <div style="background-color:#fff; border: 2px dashed #b0bec5; padding: 15px; border-radius: 8px;">
                                 <strong style="color:#546e7a;">✍️ RÁP DỮ LIỆU VÀO MẪU (DRAFTING):</strong>
                                 <br><i>(AI lấy số liệu thực tế lắp vào khung mẫu đã chọn ở trên)</i>
                                 <br>
                                 <br>👉 <b>Câu Xu hướng (Trend):</b>
                                 <br><code style="font-size: 1rem; color:#d35400; background-color:#fbe9e7; padding:4px;">[AI viết câu 1 hoàn chỉnh theo đúng Mẫu đã chọn, thay thế [...] bằng dữ liệu thật]</code>
                                 <br>
                                 <br>👉 <b>Câu Điểm nổi bật (Highlight):</b>
                                 <br><code style="font-size: 1rem; color:#d35400; background-color:#fbe9e7; padding:4px;">[AI viết câu 2 hoàn chỉnh theo đúng Mẫu đã chọn, thay thế [...] bằng dữ liệu thật]</code>
                             </div>
                         </li>

                         <!-- PHẦN 3: TỪ VỰNG -->
                         <li><b>🔑 TỪ VỰNG GỢI Ý (Vocabulary):</b>
                             <br>- <b>Verbs:</b> <i>witnessed a downward trend / saw a significant rise</i>.
                             <br>- <b>Adjectives:</b> <i>volatile (biến động) / stable (ổn định)</i>.
                             <br>- <b>Paraphrase bài này:</b> <i>[AI liệt kê 3 từ vựng sát với chủ đề bài viết]</i>
                         </li>

                         <!-- PHẦN 4: BÀI MẪU -->
                         <li><div style="background-color:#fff3e0; padding:15px; border-radius:8px; margin-top:10px; border-left: 5px solid #ff9f43;">
                             <b>📝 Nội dung mẫu (Sample Overview):</b><br>
                             <div style="margin-top:5px; font-style: italic; color: #5d4037;">
                             [AI hãy viết đoạn Overview hoàn chỉnh. <b>QUAN TRỌNG:</b> Chỉ được sử dụng đúng cấu trúc câu đã chọn ở Phần 'PHÂN TÍCH'. Lắp ghép dữ liệu vào khung mẫu.]
                             </div>
                         </div></li>
                       </ul>
    ### 🕒 PHẦN BỔ TRỢ: KHO THỜI GIAN (ĐẶT TRƯỚC BODY 1)

   - <ul>
     <!-- ================================================================================== -->
     <!-- GLOBAL MODULE: KHO THỜI GIAN VẠN NĂNG -->
     <!-- SYSTEM_ALERT: AI tham chiếu bảng này để chọn từ chỉ thời gian phù hợp cho từng phần. -->
     <!-- ================================================================================== -->
     <li>
         <div style="background-color:#e3f2fd; border:1px solid #2196f3; border-radius:5px; padding:15px; margin-bottom:15px;">
             <strong style="color:#0d47a1;">⏰ KHO THỜI GIAN VẠN NĂNG (UNIVERSAL TIME MENU):</strong>
             <br><i>(Hướng dẫn: Chọn 1 cụm từ dưới đây để điền vào biến số <b>[Time]</b> trong các Body)</i>
             
             <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin-top: 10px; font-size: 0.9rem;">
                 <!-- CỘT 1: ĐIỂM THỜI GIAN -->
                 <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #2196f3;">
                     <b style="color:#1565c0;">1. Điểm thời gian (Point):</b>
                     <ul style="margin:5px 0; padding-left:15px;">
                         <li>In <b>[Year]</b> / In the year <b>[Year]</b></li>
                         <li>At the start/beginning of the period</li>
                         <li>In the final year / By <b>[Year]</b></li>
                         <li>In <b>[Year]</b> and <b>[Year]</b> respectively</li>
                     </ul>
                 </div>

                 <!-- CỘT 2: KHOẢNG THỜI GIAN -->
                 <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #2196f3;">
                     <b style="color:#1565c0;">2. Khoảng thời gian (Duration):</b>
                     <ul style="margin:5px 0; padding-left:15px;">
                         <li>Over the <b>[Number]</b>-year period</li>
                         <li>Throughout the period</li>
                         <li>During the given period</li>
                         <li>In the subsequent years / Following this,</li>
                     </ul>
                 </div>
             </div>
         </div>
     </li>
   </ul>
                    3. **"body1_guide" (Thân bài 1 - Nhóm Nổi bật / Biến động mạnh):**
   - <ul>
     <li><b>Logic chọn nhóm (Grouping Strategy):</b> AI thực hiện: Chọn 1 hoặc 2 hạng mục (Category) có <b>số liệu cao nhất</b> hoặc <b>xu hướng thay đổi mạnh nhất/tăng trưởng rõ rệt</b> để phân tích trước.</li>
     <!-- ================================================================================== -->
     <!-- PHẦN 1: CÔNG THỨC TỔNG QUÁT (VISUAL FLOW) -->
     <!-- SYSTEM_ALERT: Hiển thị ngay đầu tiên. -->
     <!-- ================================================================================== -->
     <li>
         <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">             
             <b>► Cấu trúc chuẩn (Formula):</b> 
             <br><i>(Quy trình 3 bước chuẩn hóa cho Body 1)</i>
             <br>
             <br>
             <code style="font-size: 0.6rem; background-color: white; padding: 5px; border: 1px dashed #f9a825;">[1. So sánh Start]</code>
             <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
             <code style="font-size: 0.6rem; background-color: white; padding: 5px; border: 1px dashed #f9a825;">[2. Trend & End (Cat A)]</code>
             <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
             <code style="font-size: 0.6rem; background-color: white; padding: 5px; border: 1px dashed #f9a825;">[3. So sánh & End (Cat B)]</code>
         </div>
     </li>

     <!-- ================================================================================== -->
     <!-- SYSTEM_ALERT: 
          1. DISPLAY_MODE = RAW. Tại phần "KHO MẪU CÂU", giữ nguyên các ký tự [Variable]. KHÔNG ĐIỀN DỮ LIỆU VÀO ĐÓ.
          2. EXECUTION_MODE = STRICT. Tại phần "THỰC HÀNH", không được paraphrase. Chỉ thay thế biến số.
     <!-- ================================================================================== -->
     <li>
         <div style="background-color:#f8f9fa; border:1px solid #bdc3c7; border-radius:5px; padding:15px; margin-bottom:15px;">
         <strong style="color:#d35400;">📚 KHO MẪU CÂU BODY 1:</strong>
         <br><i>(Lưu ý: Chọn từ Menu Thời Gian để điền vào <b>[Time]</b>)</i>         
         
         <!-- KHỐI 1: CÂU MỞ ĐẦU (STARTING POINT) -->
         <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
         <br><i>(Phân tích đề bài và hình ảnh để chọn KHỐI 1 thuộc Mẫu nào?, và chỉ hiển thị công thức)</i>
             <b>► KHỐI 1: ĐIỂM XUẤT PHÁT:</b>             
                          
             <br>📍 <b>[Mẫu 1.1 - Chênh lệch rõ]</b> (A > B):
             <br><code><b>[Cat A]</b> was the dominant category, <b>significantly outstripping [Cat B]</b>, with respective figures of <b>[Data A]</b> and <b>[Data B]</b>.</code>
                          
             <br>📍 <b>[Mẫu 1.2 - Sát nút / Tương đồng (Similarity/Close)]</b> (A ≈ B):
             <br><code><b>[Cat A]</b> and <b>[Cat B]</b> started at <b>comparable levels</b> of <b>[Data A]</b> and <b>[Data B]</b> respectively.</code>
             
             <br>📍 <b>[Mẫu 1.3 - Đồng mức]</b> (A = B):
             <br><code>Both <b>[Cat A]</b> and <b>[Cat B]</b> <b>started the period at the exact same figure</b> of <b>[Data]</b>.</code>
                          
             <br>📍 <b>[Mẫu 1.4 - Trung tính]</b> (While):
             <br><code><b>[Cat A]</b> started the period at <b>[Data A]</b>, while the figure for <b>[Cat B]</b> was <b>[Data B]</b>.</code>

             <br>📍 <b>[Mẫu 1.5 - Đối chiếu]</b> (Regarding):
             <br><code><b>Regarding [Cat A]</b>, in <b>[Year 1]</b>, it stood at <b>[Data A]</b>, compared to <b>[Data B]</b> for <b>[Cat B]</b>.</code>
         </div>

         <!-- KHỐI 2: MIÊU TẢ TREND & ĐIỂM GÃY - CỰC KỲ QUAN TRỌNG -->
         <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
         <br><i>(Phân tích đề bài và hình ảnh để chọn KHỐI 2 thuộc Mẫu nào?, và chỉ hiển thị công thức)</i>
             <b>► KHỐI 2: MIÊU TẢ TREND:</b>
             
             <br>📍 <b>[Mẫu 2.1 - Before V-ing]</b> (Tăng rồi Giảm/Ngược lại):
             <br><code>The figure increased to <b>[Peak Data]</b>, <b>before falling back to finish at</b> <b>[End Data]</b>.</code>
             
             <br>📍 <b>[Mẫu 2.2 - Peak/Low & End]</b> (Đạt Đỉnh/Đáy rồi đổi chiều):
             <br>• <b>Đỉnh:</b> <code>It surged to <b>reach a peak of [Peak Data]</b> in <b>[Year]</b>, and then dropped to <b>[End Data]</b>.</code>
             <br>• <b>Đáy:</b> <code>It plunged to <b>hit a low of [Low Data]</b> in <b>[Year]</b>, before recovering to <b>[End Data]</b>.</code>
             
             <br>📍 <b>[Mẫu 2.3 - Soán ngôi & End]</b> (Vượt mặt):
             <br><code>It rose significantly, <b>surpassing [Cat B]</b> to become the leader and <b>ending at [End Data]</b>.</code>
             
             <br>📍 <b>[Mẫu 2.4 - Tăng/Giảm thẳng]</b> (Xu hướng đơn giản):
             <br><code>There was a sharp rise to <b>[End Data]</b>, making it the highest category at the end of the period.</code>
             
             <br>📍 <b>[Mẫu 2.5 - Finishing]</b> (Mệnh đề rút gọn):
             <br><code>It witnessed a steady trend, <b>finishing the period at [End Data]</b>.</code>
             
             <br>📍 <b>[Mẫu 2.6 - Dao động]</b> (Zíc-zắc & Kết thúc):
             <br><code>It showed a volatile pattern, <b>fluctuating between [Data 1] and [Data 2]</b>, before settling at <b>[End Data]</b>.</code>

             <br>📍 <b>[Mẫu 2.7 - Ổn định cao]</b> (Dùng cho đường luôn đứng nhất):
             <br><code>The figure <b>remained relatively stable</b> at <b>[Data]</b> throughout the period, maintaining its leading position.</code>

             <br>📍 <b>[Mẫu 2.8 - Plateau]</b> (Tăng/Giảm rồi đi ngang):
             <br><code>It climbed rapidly to reach <b>[Data]</b> in <b>[Year]</b>, <b>after which it leveled off/plateaued</b> for the remainder of the period.</code>

             <br>📍 <b>[Mẫu 2.9 - Recovery]</b> (Hồi phục sau khi giảm):
             <br><code>After an initial drop to <b>[Low Data]</b>, the figure <b>staged a recovery</b>, rising back to <b>[End Data]</b>.</code>
         </div>

         <!-- KHỐI 3: SO SÁNH VỚI CÁC CATEGORY CÒN LẠI -->
         <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
         <br><i>(Phân tích đề bài và hình ảnh để chọn KHỐI 3 thuộc Mẫu nào?, và chỉ hiển thị công thức)</i>
             <b>► KHỐI 3: SO SÁNH VỚI CÁC CATEGORY CÒN LẠI:</b>             
             <br>📍 <b>[Mẫu 3.1 - Tương đồng]</b>: <code>Similarly, <b>[Category B]</b> also witnessed a downward trend, <b>falling/rising to [Data]</b>.</code>
             <br>📍 <b>[Mẫu 3.2 - Đối lập]</b>: <code>In contrast, <b>[Category B]</b> <b>followed the opposite trend</b>, <b>as it increased/decreased to [Data]</b>.</code>
             <br>📍 <b>[Mẫu 3.3 - Toán học]</b>: <code><b>[Cat B]</b> followed a similar pattern, ending the period at <b>[End Data]</b>, which was <b>half/double</b> that of <b>[Cat A]</b>.</code>
             <br>📍 <b>[Mẫu 3.4 - Mức độ nhẹ hơn]</b>: <code>A similar, <b>albeit less dramatic</b>, trend was seen in <b>[Cat B]</b>, which rose/fell slightly to <b>[End Data]</b> <b>[Time]</b>.</code>
             <br>📍 <b>[Mẫu 3.5 - Cấu trúc While]</b>: <code>While <b>[Cat A]</b> [Trend A], <b>[Cat B]</b> [Trend B], finishing at <b>[End Data]</b> respectively.</code>
         </div>
         </div>
     </li>
     
     <hr style="border-top: 1px dashed #ccc; margin: 15px 0;">

     <!-- ================================================================================== -->
     <!-- PHẦN 2: GHÉP CÂU HOÀN CHỈNH -->
     <!-- SYSTEM_ALERT: QUY TRÌNH "TEMPLATE RECALL" (GỌI LẠI MẪU).
          AI BẮT BUỘC PHẢI IN RA MẪU GỐC TRƯỚC KHI ĐIỀN TỪ. -->
     <!-- ================================================================================== -->
     <li>     
         <b>✍️ THỰC HÀNH LẮP RÁP BODY 1 (CHẾ ĐỘ COPY-PASTE):</b>
         <br><i>(Quy trình: Chọn ID &rarr; In Mẫu Gốc &rarr; Khai báo &rarr; Điền vào chỗ trống)</i>
         <br>
         <br><b>✅ BƯỚC 1: Xử lý Câu mở đầu</b>
         <br>- <b>🔴 Chọn ID:</b> <i>[AI chọn Mẫu 1.X]</i>
         <br>- <b>📥 Mẫu gốc:</b> <code>[AI copy y nguyên mẫu gốc]</code>
         <br>- <b>🧩 Khai báo biến:</b>
           <br>&nbsp;&nbsp;+ [Time] = <i>(Chọn từ Menu Thời Gian)</i>
           <br>&nbsp;&nbsp;+ [Cat A] = ...
           <br>&nbsp;&nbsp;+ [Cat B] = ...
           <br>&nbsp;&nbsp;+ [Data] = ...           
         <br>- <b>📝 Điền từ:</b> <i>[AI điền dữ liệu vào mẫu]</i>
         <br>
         <br><b>✅ BƯỚC 2: Xử lý Chủ thể 1 (Trend + End)</b>
         <br>- <b>🔴 Chọn ID Trend:</b> <i>[AI chọn Mẫu 2.X]</i>
         <br>- <b>📥 Mẫu gốc:</b> <code>[AI copy y nguyên mẫu gốc]</code>
         <br>- <b>🧩 Khai báo biến:</b> [Data] = ... ; [Time] = ...
         <br>- <b>🔴 Chọn ID End:</b> <i>[AI chọn Mẫu 3.X]</i>
         <br>- <b>🧩 Khai báo biến:</b> [End Data] = ...
         <br>- <b>📝 Điền từ:</b> <i>[AI ghép 2 mẫu và điền dữ liệu]</i>
         <br>
         <br><b>✅ BƯỚC 3: Xử lý Chủ thể 2 (So sánh + Trend + End)</b>
         <br>- <b>🔴 Chọn ID So sánh:</b> <i>[AI chọn Mẫu 4.X]</i>
         <br>- <b>📥 Mẫu gốc:</b> <code>[AI copy y nguyên mẫu gốc]</code>
         <br>- <b>🧩 Khai báo biến:</b>
           <br>&nbsp;&nbsp;+ [Cat B] = ...
           <br>&nbsp;&nbsp;+ [Data] = ...
           <br>&nbsp;&nbsp;+ [Time] = <i>(Chọn từ Menu Thời Gian)</i>
         <br>- <b>📝 Điền từ:</b> <i>[AI điền dữ liệu vào mẫu]</i>
     </li>
     
     <!-- ================================================================================== -->
     <!-- PHẦN 3: BỘ CÔNG CỤ NÂNG BAND (EXAMINER'S TOOLKIT) -->
     <!-- SYSTEM_ALERT: AI KHÔNG ĐƯỢC LIỆT KÊ SƠ SÀI. 
          BẮT BUỘC PHẢI CHỌN LỌC ÍT NHẤT 02 TỪ/CỤM TỪ ĐẮT GIÁ CỦA MỖI NHÓM, PHÙ HỢP VỚI BÀI VIẾT. -->
     <!-- ================================================================================== -->
     <li>     
     <strong style="color:#d35400;">🔑 TỪ VỰNG & NGỮ PHÁP "ĂN ĐIỂM" (VOCABULARY BANK):</strong>
     <br><i>(Học sinh học cách dùng các từ này từ chính bài nháp và bài mẫu bên dưới)</i>
         <br>
         <ul style="margin-top:5px;">
            <!-- NHÓM 1: ĐỘNG TỪ TẢ XU HƯỚNG (TREND VERBS) -->
            <li><b>1. Động từ Xu hướng:</b>
                <br><i>(Thay vì chỉ dùng increase/decrease)</i>
                <br>- <b>Mạnh (Strong):</b> <i>surge / rocket</i> (Tăng vọt), <i>plunge / drop sharply</i> (Giảm mạnh).
                <br>- <b>Trải nghiệm:</b> <i>experience / undergo / witness (Chứng kiến)</i> (+ a rise/decline).
                <br><i>VD: "The figure experienced a sharp decline."</i>
                <br>- <b>Hồi phục/Rút lui:</b> <i>recover</i> (Hồi phục), <i>recede / fall back</i> (Rút xuống/Giảm lại).
            </li>

            <!-- NHÓM 2: CẤU TRÚC ĐẶC BIỆT (SPECIAL STRUCTURES) -->
            <li><b>2. Cấu trúc đặc biệt:</b>
                <br>- <b>Chủ ngữ giả:</b> <i>"Hạng mục A + <b>saw / witnessed / recorded</b> + an increase."</i>
                <br>- <b>Gấp đôi/ba:</b> <i>increased <b>twofold / threefold</b></i> (hoặc <i>doubled / tripled</i>).
                <br>- <b>Đạt đỉnh/đáy:</b> <i>reached a peak of... / reached a low of...</i>
            </li>

            <!-- NHÓM 3: TỪ VỰNG CHỦ ĐỀ CHI TIÊU -->
            <li><b>3. Topic Vocabulary (Spending):</b>
                <br><i>(Dùng khi biểu đồ nói về Tiền/Budget - Rất hay gặp)</i>
                <br>- <b>Thay cho "Spend":</b> <i>allocate (to), devote (to), apportion (to), dedicate (to).</i>
                <br>- <b>Thay cho "Budget":</b> <i>funding, investment, financial resources.</i>
                <br><i>VD: "Portugal apportioned the most to this sector."</i>
            </li>

            <!-- NHÓM 4: TRẠNG TỪ MỨC ĐỘ (ADVERBS) -->
            <li><b>4. Trạng từ mức độ (Adverbs):</b>
                <br><i>(Bắt buộc dùng trong Body 1 để tả biến động lớn)</i>
                <br>- <b>Mạnh/Nhanh:</b> <i>significantly, dramatically, sharply, rapidly.</i>
                <br>- <b>Dao động:</b> <i>fluctuated wildly</i> (dao động dữ dội) vs <i>fluctuated modestly</i> (dao động nhẹ).
            </li>
            
            <!-- NHÓM 5. GIỚI TỪ ĐƯA SỐ LIỆU -->
             <li><b>1. Giới từ đưa số liệu (Prepositions for Data):</b>
                 <br>⚠️ <i>Phân biệt chính xác:</i>
                 <br>- <b>To:</b> Tăng/giảm <b>đến</b> mức nào. <i>(rose <b>to</b> 100)</i>.
                 <br>- <b>By:</b> Tăng/giảm <b>một khoảng</b> bao nhiêu. <i>(rose <b>by</b> 10% - từ 10 lên 20)</i>.
                 <br>- <b>At:</b> Đứng <b>tại</b> mức nào. <i>(stood <b>at</b> / peak <b>at</b>)</i>.
                 <br>- <b>With:</b> Dùng trong câu mô tả kèm theo. <i>(starting <b>with</b> 10 million...)</i>.
             </li>

             <!-- 2. CẤU TRÚC CHÈN SỐ LIỆU PHỨC TẠP -->
             <li><b>2. Cấu trúc chèn số liệu (Complex Data Insertion):</b>
                 <br><i>(Không viết số liệu ngay sau động từ mãi, hãy dùng biến thể)</i>
                 <br>- <b>Reaching:</b> <i>...rose significantly, <b>reaching</b> a peak of [Data].</i>
                 <br>- <b>Of:</b> <i>...saw an increase <b>of</b> [Data] (tăng một lượng...).</i>
                 <br>- <b>Adjective + Data:</b> <i>...to a low <b>of</b> [Data] / a peak <b>of</b> [Data].</i>
             </li>

             <!-- 3. TỪ NỐI GIỮA CÁC CÂU (LINKING DEVICES) -->
             <li><b>3. Từ nối chuyển mạch (Linking Devices):</b>
                 <br><i>(Dùng để đầu câu, giúp đoạn văn liền mạch)</i>
                 <br>- <b>Thời gian (Time markers):</b> <i>At the start of the period, ... / Thereafter, ... / In the subsequent years, ... / In the final year, ...</i>
                 <br>- <b>So sánh (Comparison):</b> <i>In contrast, ... / By contrast, ... / Similarly, ... / Likewise, ...</i>
                 <br>- <b>Chuyển ý:</b> <i>Regarding [Category A], ... / Turning to [Category B], ...</i>
             </li>
         </ul>
     </li>

     <!-- ================================================================================== -->
     <!-- PHẦN 4: GHÉP CÂU HOÀN CHỈNH -->
     <!-- ================================================================================== -->
     <li><div style="background-color:#fff8e1; padding:15px; border-radius:8px; margin-top:10px; border-left: 5px solid #ffa502;">
         <b>📝 KẾT QUẢ BODY 1 (FINAL OUTPUT):</b><br>
         <div style="margin-top:5px; font-style: italic; color: #5d4037;">
         [AI nối các câu ở phần 'Điền từ' lại thành đoạn văn.]
         <b>Checklist kiểm tra:</b><br>
         1. Có câu so sánh năm đầu không?<br>
         2. Có dùng cấu trúc "Before V-ing" hoặc "Reach a peak" không?<br>
         3. Có từ vựng trong danh sách trên không?]
         </div>
     </div></li>
   </ul>

                    4. **"body2_guide" (Thân bài 2 - Quy trình Chọn & Điền):**
   - <ul>
     <li><b>Logic chọn nhóm (Grouping Strategy):</b> AI thực hiện: Phân tích tất cả các hạng mục (Category) chưa được nhắc đến ở Body 1. Thường là nhóm có xu hướng <b>NGƯỢC LẠI</b> (Giảm/Dao động) hoặc số liệu <b>THẤP HƠN</b>.</li>
     <!-- ================================================================================== -->     
     <!-- PHẦN 1: CÔNG THỨC TỔNG QUÁT -->
     <!-- SYSTEM_ALERT: 
          1. DISPLAY_MODE = RAW. Tại phần "KHO MẪU CÂU", giữ nguyên các ký tự [Variable]. KHÔNG ĐIỀN DỮ LIỆU VÀO ĐÓ.
          2. EXECUTION_MODE = STRICT. Tại phần "THỰC HÀNH", không được paraphrase. Chỉ thay thế biến số. -->
     <!-- ================================================================================== -->
     <li>
         <div style="background-color:#f8f9fa; border:1px solid #bdc3c7; border-radius:5px; padding:15px; margin-bottom:15px;">
         <strong style="color:#2980b9;">📚 KHO MẪU CÂU BODY 2:</strong>         
         <br><i>(Lưu ý: Chọn từ Menu Thời Gian để điền vào <b>[Time]</b>)</i>
         
         <!-- KHỐI 1: CÂU CHUYỂN ĐOẠN -->
         <div style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; margin-top:5px;">
         <br><i>(Phân tích đề bài và hình ảnh để chọn KHỐI 1 thuộc Mẫu nào?, và chỉ hiển thị công thức)</i>
             <b>► KHỐI 1: CHUYỂN ĐOẠN:</b>
             <br>📍 <b>[Mẫu 1.1 - Trung tính]</b> (Dùng để liệt kê tiếp):
             <br><code>Turning to the remaining categories (<b>[Cat C]</b> and <b>[Cat D]</b>), the figure for <b>[Cat C]</b> started at <b>[Data]</b> <b>[Time]</b>.</code>
             
             <br>📍 <b>[Mẫu 1.2 - Đối lập]</b> (Dùng khi xu hướng ngược Body 1):
             <br><code><b>In contrast</b>, a completely different trend was observed in <b>[Cat C]</b>, which began the period at <b>[Data]</b> <b>[Time]</b>.</code>
                          
             <br>📍 <b>[Mẫu 1.3 - Song song]</b> (Dùng khi diễn ra cùng lúc):
             <br><code><b>Meanwhile</b>, the figure for <b>[Cat C]</b> stood at <b>[Data]</b> <b>[Time]</b>.</code>             
         </div>

         <!-- KHỐI 2: DIỄN BIẾN & KẾT THÚC (GỘP) -->
         <div style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; margin-top:5px;">
         <br><i>(Phân tích đề bài và hình ảnh để chọn KHỐI 2 thuộc Mẫu nào?, và chỉ hiển thị công thức)</i>
             <b>► KHỐI 2: DIỄN BIẾN & KẾT THÚC:</b>
                          
             <br>📍 <b>[Mẫu 2.1 - Dao động]</b> (Zíc-zắc):             
             <br><code>It <b>fluctuated (wildly/moderately)</b> around <b>[Avg Data]</b>, before settling at <b>[End Data]</b> <b>[Time]</b>.</code>
             
             <br>📍 <b>[Mẫu 2.2 - Ổn định]</b> (Đi ngang):             
             <br><code>The figure <b>remained relatively stable</b> at <b>[Data]</b> <b>[Time]</b>.</code>

             <br>📍 <b>[Mẫu 2.3 - Tăng/Giảm chuẩn]</b> (Dùng cho cả Mạnh và Nhẹ):             
             <br><code>It witnessed a <b>[Adjective]</b> <b>(increase/decrease)</b>, finishing the period at <b>[End Data]</b>.</code>
         </div>
        
         <!-- KHỐI 3: SO SÁNH NỘI BỘ -->
         <div style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; margin-top:5px;">
         <br><i>(Phân tích đề bài và hình ảnh để chọn KHỐI 3 thuộc Mẫu nào?, và chỉ hiển thị công thức)</i>
             <b>► KHỐI 3: SO SÁNH NHÓM CÒN LẠI:</b>
             
             <br>📍 <b>[Mẫu 3.1 - Tương đồng]</b>:
             <br><code>A <b>similar pattern</b> was seen in <b>[Cat D]</b>, which rose/fell to <b>[Data]</b> <b>[Time]</b>.</code>
             
             <br>📍 <b>[Mẫu 3.2 - Gấp lần]</b> (Toán học):
             <br><code>The figure for <b>[Cat C]</b> was <b>(double / two times as high as)</b> that of <b>[Cat D]</b> <b>[Time]</b>.</code>
         </div>
         </div>
     </li>
     
     <hr style="border-top: 1px dashed #ccc; margin: 15px 0;">

     <!-- ================================================================================== -->
     <!-- PHẦN 2: THỰC HÀNH TƯ DUY & LẮP RÁP -->
     <!-- ================================================================================== -->
     <li>     
         <b>✍️ THỰC HÀNH LẮP RÁP BODY 2 (STEP-BY-STEP):</b>
         <br>
         <br><b>✅ BƯỚC 1: Xử lý Câu chuyển đoạn</b>
         <br>- <b>🔴 Chọn ID:</b> <i>[AI chọn Mẫu 1.X]</i>
         <br>- <b>📥 Mẫu gốc:</b> <code>[AI copy y nguyên mẫu gốc]</code>
         <br>- <b>🧩 Khai báo biến:</b>
           <br>&nbsp;&nbsp;+ [Cat C] = ...
           <br>&nbsp;&nbsp;+ [Cat D] (nếu có) = ...
           <br>&nbsp;&nbsp;+ [Data] = ...
           <br>&nbsp;&nbsp;+ [Time] = <i>(Chọn từ Menu Thời Gian)</i>
         <br>- <b>📝 Kết quả ghép:</b> <i>[AI thay biến vào mẫu]</i>
         <br>
         <br><b>✅ BƯỚC 2: Xử lý Diễn biến & Kết thúc</b>
         <br>- <b>🔴 Chọn ID:</b> <i>[AI chọn Mẫu 2.X]</i>
         <br>- <b>📥 Mẫu gốc:</b> <code>[AI copy y nguyên mẫu gốc]</code>
         <br>- <b>🧩 Khai báo biến:</b>
           <br>&nbsp;&nbsp;+ [Avg/End Data] = ...
           <br>&nbsp;&nbsp;+ [Adjective] (nếu có) = ...
           <br>&nbsp;&nbsp;+ [Time] = <i>(Chọn từ Menu Thời Gian)</i>
         <br>- <b>📝 Kết quả ghép:</b> <i>[AI thay biến vào mẫu]</i>
         <br>
         <br><b>✅ BƯỚC 3: Xử lý So sánh nội bộ</b>
         <br>- <b>🔴 Chọn ID:</b> <i>[AI chọn Mẫu 3.X]</i>
         <br>- <b>📥 Mẫu gốc:</b> <code>[AI copy y nguyên mẫu gốc]</code>
         <br>- <b>🧩 Khai báo biến:</b>
           <br>&nbsp;&nbsp;+ [Cat C] = ...
           <br>&nbsp;&nbsp;+ [Cat D] = ...
           <br>&nbsp;&nbsp;+ [Data] = ...
           <br>&nbsp;&nbsp;+ [Time] = <i>(Chọn từ Menu Thời Gian)</i>
         <br>- <b>📝 Kết quả ghép:</b> <i>[AI thay biến vào mẫu]</i>
     </li>      

     <!-- ================================================================================== -->
     <!-- PHẦN 3: GIẢI MÃ TỪ VỰNG "ĂN ĐIỂM" TRONG BÀI (DECODING VOCABULARY) -->
     <!-- SYSTEM_ALERT: AI KHÔNG ĐƯỢC LIỆT KÊ SƠ SÀI. 
          BẮT BUỘC PHẢI CHỌN LỌC ÍT NHẤT 02 TỪ/CỤM TỪ ĐẮT GIÁ CỦA MỖI NHÓM, PHÙ HỢP VỚI BÀI VIẾT. -->
     <!-- ================================================================================== -->
     <li>
         <strong style="color:#d35400;">🔑 TỪ VỰNG & NGỮ PHÁP "ĂN ĐIỂM" (VOCABULARY BANK):</strong>
         <br><i>(Học sinh học cách dùng các từ này từ chính bài nháp và bài mẫu bên dưới)</i>
         
         <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 15px; margin-top: 10px; background-color:#fffcf5; padding:15px; border-radius:8px; border: 1px solid #f39c12;">
            <!-- CỘT 1: LIÊN KẾT & XU HƯỚNG -->
            <div>
                <b style="color:#e67e22;">1. Liên kết & Chuyển động (Logic & Trends):</b>
                <ul style="margin:5px 0; padding-left:20px; font-size:0.9rem;">
                    <li><b>Từ nối chuyển đoạn:</b> <i>[AI trích cụm đã dùng, VD: Turning to...]</i></li>
                    <li><b>Từ nối giữa câu:</b> <i>[AI trích cụm đã dùng, VD: Subsequently...]</i></li>
                    <li><b>Cụm động từ:</b> <i>[AI trích cụm đã dùng, VD: showed a volatile pattern]</i>: [Nghĩa TV]</li>
                    <li><b>Cấu trúc ổn định:</b> <i>[AI trích cụm đã dùng, VD: remained constant]</i>: [Nghĩa TV]</li>
                </ul>
            </div>

            <!-- CỘT 2: ĐỘ CHÍNH XÁC & NGỮ PHÁP -->
            <div>
                <b style="color:#e67e22;">2. Độ chính xác & Ngữ pháp (Maths & Grammar):</b>
                <ul style="margin:5px 0; padding-left:20px; font-size:0.9rem;">
                    <li><b>So sánh gấp lần:</b> <i>[AI trích cụm, VD: twice as high as]</i></li>
                    <li><b>Trạng từ mức độ:</b> <i>[AI trích, VD: marginally / significantly]</i></li>
                    <li><b>Kỹ thuật đưa số liệu:</b> <i>[AI trích, VD: ending the period at]</i></li>
                    <li><b>Cấu trúc câu phức:</b> <i>[AI trích cấu trúc After V-ing / which was...]</i></li>
                </ul>
            </div>
         </div>
     </li>

     <!-- ================================================================================== -->
     <!-- PHẦN 4: GHÉP CÂU HOÀN CHỈNH -->
     <!-- ================================================================================== -->
     <li><div style="background-color:#e1f5fe; padding:15px; border-radius:8px; margin-top:10px; border-left: 5px solid #03a9f4;">
         <b>📝 KẾT QUẢ BODY 2 (FINAL OUTPUT):</b><br>
         <div style="margin-top:5px; font-style: italic; color: #5d4037;">
         [AI nối các câu ở phần 'Kết quả ghép' lại thành đoạn văn.]
         </div>
     </div></li>
   </ul>

                    # =================================================================
                    # 🟡 TRƯỜNG HỢP 3: CÁC DẠNG KHÁC (MAP, PROCESS, MIXED)
                    # =================================================================
                    *(Tự động điều chỉnh hướng dẫn phù hợp với đặc thù từng dạng).*

                    **YÊU CẦU TRÌNH BÀY:**
                    - Dùng thẻ HTML `<ul>`, `<li>`, `<b>`, `<i>`, `<code style='color:#d63384'>` để highlight.
                    - Giải thích ngắn gọn, dễ hiểu.

                    **JSON OUTPUT FORMAT:**
                    {
                        "task_type": "Tên loại bài (Ví dụ: Static Pie Charts)",
                        "intro_guide": "HTML string...",
                        "overview_guide": "HTML string...",
                        "body1_guide": "HTML string...",
                        "body2_guide": "HTML string..."
                    }
                    """
                    
                    # Gọi AI
                    res, _ = generate_content_with_failover(prompt_guide + "\nĐề bài: " + question_input, img_data, json_mode=True)
                    if res:
                        data = parse_guide_response(res.text)
                    # Dù AI trả về gì, ta cũng phải gán guide_data để App không bị kẹt ở Step 1
                        st.session_state.guide_data = data if data else {
                            "task_type": "Task 1", "intro_guide": "AI Error - Please try again", 
                            "overview_guide": "", "body1_guide": "", "body2_guide": ""
                    }
                    st.session_state.step = 2
                    st.rerun() # Buộc Streamlit vẽ lại giao diện Phase 2 ngay lập tức

# ==========================================
# 6. UI: PHASE 2 - WRITING PRACTICE (ULTIMATE STICKY)
# ==========================================
if st.session_state.step == 2 and st.session_state.guide_data:
    
    # --- 1. CSS "ĐÓNG BĂNG" CỘT TRÁI ---
    st.markdown("""
        <style>
            /* Nhắm vào container chứa cả 2 cột */
            [data-testid="stHorizontalBlock"] {
                align-items: flex-start !important;
            }

            /* Nhắm vào cột đầu tiên (Cột Trái) */
            [data-testid="stHorizontalBlock"] > div:nth-child(1) {
                position: -webkit-sticky !important;
                position: sticky !important;
                top: 2rem !important;
                z-index: 999 !important;
            }

            /* Cố định chiều cao vùng hiển thị đề bài để không bị trôi */
            [data-testid="stHorizontalBlock"] > div:nth-child(1) > div:nth-child(1) {
                max-height: 95vh !important;
                overflow-y: auto !important;
                padding-right: 10px !important;
            }

            /* Tùy chỉnh thanh cuộn cho cột trái */
            [data-testid="stHorizontalBlock"] > div:nth-child(1) > div:nth-child(1)::-webkit-scrollbar {
                width: 4px;
            }
            [data-testid="stHorizontalBlock"] > div:nth-child(1) > div:nth-child(1)::-webkit-scrollbar-thumb {
                background: #cccccc;
                border-radius: 10px;
            }
            
            /* Tăng khoảng cách giữa các ô nhập liệu bên phải */
            .stTextArea {
                margin-bottom: 1rem !important;
            }
        </style>
    """, unsafe_allow_html=True)

    data = st.session_state.guide_data

    # --- 2. HÀM RENDER (ĐÃ SỬA LỖI EXPANDER & HTML) ---
    def render_writing_section(title, guide_key, input_key):
        # Tính số từ hiện tại
        current_text = st.session_state.get(input_key, "")
        word_count = len(current_text.split())
        
        # Tiêu đề + Số từ
        c_title, c_count = st.columns([8, 2])
        with c_title:
            st.markdown(f"#### {title}")
        with c_count:
            st.markdown(f"""
            <div style="text-align: right; padding-top: 10px;">
                <span style="background-color: #F1F5F9; color: #64748B; padding: 4px 10px; border-radius: 12px; font-size: 12px; border: 1px solid #E2E8F0;">
                    {word_count} words
                </span>
            </div>
            """, unsafe_allow_html=True)

        # Phần hướng dẫn (FIX: Markdown thụt vào trong 'with')
        with st.expander(f"💡 Hướng dẫn viết {title}", expanded=False):
            g_text = data.get(guide_key, "Không có hướng dẫn.")
            
            # FIX: Xóa xuống dòng để không bị hiểu lầm là Code Block
            if g_text:
                import re
                g_text = re.sub(r'\n\s*', ' ', g_text)
                
            st.markdown(f"<div class='guide-box'>{g_text}</div>", unsafe_allow_html=True)
        
        # Ô nhập liệu (Nằm ngoài 'with', trả về kết quả)
        return st.text_area(label=title, height=200, key=input_key, placeholder=f"Bắt đầu viết {title} tại đây...", label_visibility="collapsed")

    # --- 3. CHIA CỘT LAYOUT (4-6) ---
    col_left, col_right = st.columns([4, 6], gap="large")

    with col_left:
        st.subheader("📄 Đề bài & Hình ảnh")
        # Khung chứa đề bài
        st.markdown(f"""
            <div style="background-color: #F1F5F9; padding: 20px; border-radius: 10px; border: 1px solid #CBD5E1; line-height: 1.6; color: #1E293B; margin-bottom: 15px;">
                <b>Question:</b><br><i>{st.session_state.saved_topic}</i>
            </div>
        """, unsafe_allow_html=True)
        
        # Hình ảnh biểu đồ
        if st.session_state.saved_img:
            st.image(st.session_state.saved_img, width="stretch")
        
        st.info(f"📌 **Dạng bài:** {data.get('task_type')}")

    with col_right:
        st.subheader("✍️ Khu vực viết bài")
        
        # --- ĐÃ XÓA WORD COUNT TỔNG Ở ĐÂY ---

        # Render các ô nhập liệu (Có word count riêng từng phần)
        intro_text = render_writing_section("Introduction", "intro_guide", "in_intro")
        overview_text = render_writing_section("Overview", "overview_guide", "in_overview")
        body1_text = render_writing_section("Body 1", "body1_guide", "in_body1")
        body2_text = render_writing_section("Body 2", "body2_guide", "in_body2")

        st.markdown("---")
        
        # --- TÍNH TỔNG SỐ TỪ VÀ HIỂN THỊ Ở DƯỚI CÙNG ---
        def count_w(k): return len(st.session_state.get(k, "").split())
        total_wc = count_w("in_intro") + count_w("in_overview") + count_w("in_body1") + count_w("in_body2")
        
        # Hiển thị Total Word Count
        st.markdown(f"""
            <div style="text-align: right; margin-bottom: 15px;">
                <span style="background-color: #10B981; color: white; padding: 8px 20px; border-radius: 20px; font-weight: bold; font-size: 16px; box-shadow: 0 2px 5px rgba(0,0,0,0.1);">
                    Total word count: {total_wc}
                </span>
            </div>
        """, unsafe_allow_html=True)

        # Nút chấm điểm
        if st.button("🎓 Gửi bài chấm điểm (Examiner Pro)", type="primary", width="stretch"):
            if total_wc < 30:
                st.warning("⚠️ Bài viết quá ngắn, AI không thể chấm điểm chính xác.")
            else:
                with st.status("👨‍🏫 Giám khảo đang chấm bài...") as status:
                    total_essay = f"{intro_text}\n\n{overview_text}\n\n{body1_text}\n\n{body2_text}".strip()
                    # Sử dụng biến saved_topic để tránh lỗi NameError
                    prompt_grade = GRADING_PROMPT_TEMPLATE.replace('{{TOPIC}}', st.session_state.saved_topic).replace('{{ESSAY}}', total_essay)
                    
                    res_grade, _ = generate_content_with_failover(prompt_grade, st.session_state.saved_img, json_mode=False)
                    
                    if res_grade:
                        # process_grading_response là hàm bóc tách Text và JSON bạn đã có
                        mk_text, p_data = process_grading_response(res_grade.text)
                        st.session_state.grading_result = {
                            "data": p_data, "markdown": mk_text,
                            "essay": total_essay, "topic": st.session_state.saved_topic
                        }
                        st.session_state.step = 3
                        status.update(label="✅ Đã chấm xong!", state="complete", expanded=False)
                        st.rerun()
                    else:
                        status.update(label="❌ Lỗi kết nối AI", state="error")

# ==========================================
# 7. UI: PHASE 3 - GRADING RESULT (FINAL POLISHED)
# ==========================================
if st.session_state.step == 3 and st.session_state.grading_result:
    
    # --- 1. CSS TINH CHỈNH CUỐI CÙNG ---
    st.markdown("""
        <style>
            /* 1. Layout 2 cột */
            [data-testid="stHorizontalBlock"] {
                align-items: flex-start !important;
            }

            /* 2. Style cho 2 cái Hộp lớn (Container) */
            /* Streamlit tự tạo container có viền, ta chỉ cần chỉnh background app cho nổi bật */
            .stApp {
                background-color: #ffffff;
            }

            /* 3. Style Bài viết: Tự động xuống dòng, không cuộn ngang */
            .essay-review-box {
                background-color: #f8fafc;
                border: 1px solid #cbd5e1;
                border-radius: 6px;
                padding: 15px; /* Tăng padding cho dễ đọc */
                
                font-family: 'Inter', sans-serif;
                font-size: 0.95rem;
                line-height: 1.6;
                color: #334155;
                
                /* QUAN TRỌNG: Ép xuống dòng */
                white-space: pre-wrap !important;       /* Giữ dòng mới nhưng wrap text */
                word-wrap: break-word !important;       /* Ngắt từ dài */
                overflow-wrap: break-word !important;   /* Hỗ trợ trình duyệt hiện đại */
                max-width: 100%;                        /* Không vượt quá chiều rộng hộp cha */
            }

            /* 4. Thanh cuộn đẹp */
            ::-webkit-scrollbar { width: 6px; height: 6px; }
            ::-webkit-scrollbar-thumb { background: #cbd5e1; border-radius: 4px; }
        </style>
    """, unsafe_allow_html=True)

    res = st.session_state.grading_result
    g_data = res["data"]
    analysis_text = res["markdown"]
    
    # --- 2. CHIA CỘT (Không cần tiêu đề to nữa) ---
    c1, c2 = st.columns([4, 6], gap="medium")

    # === HỘP TRÁI: THÔNG TIN ĐỐI CHIẾU ===
    with c1:
        # Hộp chứa có chiều cao cố định để tạo thanh cuộn
        with st.container(height=750, border=True):
            st.markdown("#### 📄 Thông tin đối chiếu")
            
            # Ảnh
            if st.session_state.saved_img:
                st.image(st.session_state.saved_img, width="stretch")
            
            st.markdown("---")
            
            # Đề bài
            with st.expander("📌 Đề bài (Prompt)", expanded=False):
                st.info(st.session_state.saved_topic)
                
            # Bài viết (Đã áp dụng class mới để không tràn)
            st.markdown("**✍️ Bài viết của bạn:**")
            st.markdown(f'<div class="essay-review-box">{html.escape(res["essay"])}</div>', unsafe_allow_html=True)

    # === HỘP PHẢI: KẾT QUẢ CHẤM ===
    with c2:
        with st.container(height=750, border=True):
            st.markdown("#### 👨‍🏫 Examiner Analysis")
            
            # Bảng điểm
            scores = g_data.get("originalScore", {})
            st.markdown(f"""
            <div style="background-color: #ecfdf5; border: 1px solid #6ee7b7; border-radius: 10px; padding: 15px; display: flex; align-items: center; justify-content: space-between; margin-bottom: 20px;">
                <div style="text-align: center;">
                    <span style="color: #047857; font-weight: bold; font-size: 0.9rem;">BAND SCORE</span><br>
                    <span style="color: #059669; font-weight: 900; font-size: 2.5rem; line-height: 1;">{scores.get("overall", "-")}</span>
                </div>
                <div style="display: flex; gap: 15px; text-align: center;">
                    <div><small style="color:#047857;">TA</small><br><b style="color:#059669; font-size:1.1rem;">{scores.get("task_achievement", "-")}</b></div>
                    <div><small style="color:#047857;">CC</small><br><b style="color:#059669; font-size:1.1rem;">{scores.get("cohesion_coherence", "-")}</b></div>
                    <div><small style="color:#047857;">LR</small><br><b style="color:#059669; font-size:1.1rem;">{scores.get("lexical_resource", "-")}</b></div>
                    <div><small style="color:#047857;">GRA</small><br><b style="color:#059669; font-size:1.1rem;">{scores.get("grammatical_range", "-")}</b></div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            # Tabs chi tiết
            tab1, tab2, tab3, tab4 = st.tabs(["📝 Phân tích chuyên sâu", "🔴 Lỗi Ngữ pháp và Từ vựng", "🔵 Lỗi Mạch lạc", "✍️ Bài sửa"])
            
            with tab1:
                st.markdown(analysis_text if analysis_text and len(analysis_text) > 50 else "Chưa có dữ liệu phân tích.")

            with tab2:
                micro = [e for e in g_data.get('errors', []) if e.get('category') in ['Grammar', 'Vocabulary', 'Ngữ pháp', 'Từ vựng']]
                if not micro: st.success("✅ Tuyệt vời! Không có lỗi ngữ pháp lớn.")
                for i, err in enumerate(micro):
                    badge = "#DCFCE7" if err.get('category') in ['Grammar','Ngữ pháp'] else "#FEF9C3"
                    st.markdown(f"""
                    <div class="error-card">
                        <b>#{i+1} {err.get('type')}</b>
                        <div style="background:{badge}; padding:5px; border-radius:4px; margin:5px 0;">
                            <s>{err.get('original')}</s> ➔ <b>{err.get('correction')}</b>
                        </div>
                        <small><i>{err.get('explanation')}</i></small>
                    </div>
                    """, unsafe_allow_html=True)

            # Tab 3: Lỗi Mạch lạc (Macro) - ĐÃ SỬA LỖI HIỂN THỊ RAW CODE
            with tab3:
                macro = [e for e in g_data.get('errors', []) if e.get('category') not in ['Grammar', 'Vocabulary', 'Ngữ pháp', 'Từ vựng']]
                if not macro: 
                    st.success("✅ Cấu trúc tốt.")
                for err in macro:
                    # Lưu ý: Các thẻ HTML bên dưới được viết sát lề trái của chuỗi f-string
                    # để tránh bị Markdown hiểu nhầm là Code Block.
                    st.markdown(f"""
<div class="error-card-container" style="border-left: 4px solid #3b82f6;">
    <div style="font-weight:bold; color:#1e40af; margin-bottom:5px;">{err.get('type')}</div>
    <div style="background-color:#eff6ff; padding:8px; border-radius:4px; margin-bottom:8px; border:1px dashed #93c5fd;">
        <span style="font-size:0.8rem; font-weight:bold; color:#60a5fa;">TRÍCH DẪN:</span><br>
        <span style="font-family:monospace; color:#1e3a8a;">"{err.get('original', 'N/A')}"</span>
    </div>
    <div style="margin-bottom:5px;"><b>Vấn đề:</b> {err.get('explanation')}</div>
    <div style="color:#059669;"><b>👉 Gợi ý:</b> {err.get('correction')}</div>
</div>
""", unsafe_allow_html=True)

            with tab4:
                st.markdown(f'<div class="annotated-text">{g_data.get("annotatedEssay", "")}</div>', unsafe_allow_html=True)

            st.markdown("---")
            
            # Download & Reset
            d1, d2 = st.columns(2)
            docx = create_docx(g_data, res['topic'], res['essay'], analysis_text)
            d1.download_button("📥 Tải báo cáo (.docx)", docx, "IELTS_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            if st.button("🔄 Làm bài mới (Reset)", width="stretch"):
                for k in ["step", "guide_data", "grading_result", "saved_topic", "saved_img"]: st.session_state[k] = None
                st.session_state.step = 1
                st.rerun()
# ==========================================
# FOOTER (HIỂN THỊ Ở MỌI STEP)
# ==========================================
st.markdown("""
    <style>
        .footer-text {
            text-align: center; 
            color: #94a3b8; 
            font-size: 0.8rem; 
            font-family: 'Inter', sans-serif; 
            padding-top: 15px;      /* Giảm đệm trên */
            padding-bottom: 0px;   /* Giảm đệm dưới */
            border-top: 1px solid #e2e8f0;
            margin-top: 30px;       /* Khoảng cách với nội dung bên trên */
        }
    </style>
    <div class="footer-text">
        © 2025 Developed by <b>Albert Nguyen</b>
    </div>
""", unsafe_allow_html=True)
